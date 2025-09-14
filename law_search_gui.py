import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog
import requests
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
import PyPDF2
import google.generativeai as genai
import re
from dotenv import load_dotenv
import os

load_dotenv()
api_key = os.environ.get('GEMINI_API_KEY')
genai.configure(api_key=api_key)

class LawSearchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("광역지자체 조례 검색, 비교, 분석 프로그램")
        self.root.geometry("1000x800")
        
        # API 설정
        self.OC = "climsneys85"  # 이메일 ID
        self.search_url = "http://www.law.go.kr/DRF/lawSearch.do"
        self.detail_url = "http://www.law.go.kr/DRF/lawService.do"
        
        # 광역지자체 코드 및 이름
        self.metropolitan_govs = {
            '6110000': '서울특별시',
            '6260000': '부산광역시',
            '6270000': '대구광역시',
            '6280000': '인천광역시',
            '6290000': '광주광역시',
            '6300000': '대전광역시',
            '6310000': '울산광역시',
            '6360000': '세종특별자치시',
            '6410000': '경기도',
            '6420000': '강원도',
            '6430000': '충청북도',
            '6440000': '충청남도',
            '6450000': '전라북도',
            '6460000': '전라남도',
            '6470000': '경상북도',
            '6480000': '경상남도',
            '6500000': '제주특별자치도'
        }
        
        self.create_widgets()
        
    def create_widgets(self):
        # 스타일 설정 (글래스모피즘+자주/핑크 계열)
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background='#f7f9fa')
        style.configure('TLabel', background='#f7f9fa', font=("맑은 고딕", 11))
        # 자주/핑크 계열 버튼
        style.configure('PurpleGlass.TButton', font=("맑은 고딕", 12, "bold"), padding=(24, 2, 6, 2), background='#d946ef', foreground='white', borderwidth=0, relief='flat')
        style.map('PurpleGlass.TButton', background=[('active', '#e75480')], foreground=[('active', 'white')])
        style.configure('Search.TButton', font=("맑은 고딕", 11, "bold"), padding=8, background='#f472b6', foreground='white', borderwidth=0, relief='flat')
        style.map('Search.TButton', background=[('active', '#d946ef')], foreground=[('active', 'white')])
        style.configure('Order.TLabel', font=("맑은 고딕", 12, "bold"), background='#f7f9fa', foreground='#a21caf')
        style.configure('Result.TLabel', font=("맑은 고딕", 11, "bold"), background='#f7f9fa', foreground='#d946ef')

        self.root.configure(bg='#f7f9fa')

        # 상단 검색 프레임 (조례명 입력 + 검색)
        search_frame = ttk.Frame(self.root, padding="10")
        search_frame.pack(fill=tk.X)
        ttk.Label(search_frame, text="조례명:").pack(side=tk.LEFT)
        self.search_entry = ttk.Entry(search_frame, width=40, font=("맑은 고딕", 11))
        self.search_entry.pack(side=tk.LEFT, padx=5)
        search_button = ttk.Button(search_frame, text="검색", command=self.search_laws, style='Search.TButton')
        search_button.pack(side=tk.LEFT, padx=5)

        # 순서 안내 프레임
        order_frame = ttk.Frame(self.root, padding="5")
        order_frame.pack(fill=tk.X)
        ttk.Label(order_frame, text="[순서]", style='Order.TLabel').pack(anchor=tk.W, pady=(0, 4))

        # 1단계 버튼
        save_button = ttk.Button(self.root, text="1. 검색된 타 시도 조례를 3단비교 형태로 MS워드 저장", command=self.save_selected_to_word, style='PurpleGlass.TButton')
        save_button.pack(fill=tk.X, padx=8, pady=(0, 16), ipady=2)
        # 2단계 버튼
        upload_pdf_button = ttk.Button(self.root, text="2. 제정, 개정할 조례안 PDF파일 업로드", command=self.upload_pdf, style='PurpleGlass.TButton')
        upload_pdf_button.pack(fill=tk.X, padx=8, pady=(0, 16), ipady=2)
        # 3단계 버튼
        compare_button = ttk.Button(self.root, text="3. 제개정 조례안과 타시도 조례안 비교 분석한 후 MS워드로 저장", command=self.debug_compare_pdf_with_laws, style='PurpleGlass.TButton')
        compare_button.pack(fill=tk.X, padx=8, pady=(0, 16), ipady=2)

        # 검색결과 안내문
        result_label = ttk.Label(self.root, text="[검색결과]", style='Result.TLabel')
        result_label.pack(anchor=tk.W, padx=10, pady=(12,0))

        # 결과 표시 영역
        result_frame = ttk.Frame(self.root, padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True)
        self.result_text = scrolledtext.ScrolledText(result_frame, wrap=tk.WORD, width=100, height=40, font=("맑은 고딕", 11), background='#ffffff', foreground='#222', borderwidth=2, relief='groove')
        self.result_text.pack(fill=tk.BOTH, expand=True)
        self.result_text.tag_configure("red_bold", foreground="#e74c3c", font=("맑은 고딕", 10, "bold"))
        
    def get_ordinance_detail(self, ordinance_id):
        params = {
            'OC': self.OC,
            'target': 'ordin',
            'ID': ordinance_id,
            'type': 'XML'
        }
        try:
            response = requests.get(self.detail_url, params=params)
            root = ET.fromstring(response.text)
            articles = []
            for article in root.findall('.//조'):
                content = article.find('조내용').text if article.find('조내용') is not None else ""
                # 태그 제거
                if content:
                    content = content.replace('<![CDATA[', '').replace(']]>', '')
                    content = content.replace('<p>', '').replace('</p>', '\n')
                    content = content.replace('<br/>', '\n')
                    content = content.replace('<br>', '\n')
                    content = content.replace('&nbsp;', ' ')
                    content = content.strip()
                if content:
                    articles.append(content)
            return articles
        except Exception:
            return []
        
    def search_laws(self):
        search_query = self.search_entry.get().strip()
        if not search_query:
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, "검색어를 입력해주세요.")
            return
        self.result_text.delete(1.0, tk.END)  # 검색 시작 전에 한 번만!
        total_count = 0
        law_names = []
        all_articles = []
        # 각 광역지자체별로 검색
        for org_code, metro_name in self.metropolitan_govs.items():
            params = {
                'OC': self.OC,
                'target': 'ordin',
                'type': 'XML',
                'query': search_query,
                'display': 100,
                'search': 1,  # 제목만 검색
                'sort': 'ddes',
                'page': 1,
                'org': org_code
            }
            try:
                response = requests.get(self.search_url, params=params)
                root = ET.fromstring(response.text)
                total_laws = len(root.findall('.//law'))
                if total_laws > 0:
                    self.result_text.insert(tk.END, f"\n{metro_name} 검색 결과: {total_laws}건\n")
                for law in root.findall('.//law'):
                    ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                    ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                    기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                    if 기관명 != metro_name:
                        continue  # 본청이 아니면 건너뜀
                    # 검색어 매칭 로직 - 검색어의 모든 단어가 조례명에 포함되어야 함 (공백/대소문자 무시)
                    search_terms = [term.lower() for term in search_query.split() if term.strip()]
                    ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                    if not all(term in ordinance_name_clean for term in search_terms):
                        continue
                    total_count += 1
                    # 조례명을 빨간색 굵게 표시
                    self.result_text.insert(tk.END, ordinance_name + "\n", "red_bold")
                    articles = self.get_ordinance_detail(ordinance_id)
                    law_names.append(ordinance_name)
                    all_articles.append(articles)
                    if articles:
                        for content in articles:
                            self.result_text.insert(tk.END, f"{content}\n\n")
                    else:
                        self.result_text.insert(tk.END, "  (조문 없음)\n")
                    self.result_text.insert(tk.END, "\n")
            except Exception as e:
                self.result_text.insert(tk.END, f"{metro_name} 검색 중 오류 발생: {str(e)}\n")
                continue
        self.result_text.insert(tk.END, f"전체 검색 결과: {total_count}건 (자치법규)\n\n")
        # 본청 조례명/조문 리스트 저장
        self.last_search_law_names = law_names
        self.last_search_all_articles = all_articles

    def save_to_word(self, law_names, all_articles):
        from datetime import datetime
        from docx.shared import Mm, RGBColor
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt
        doc = Document()
        
        # 용지 방향을 가로, 크기는 A3로
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(420)
        section.page_height = Mm(297)
        
        # 조례를 3개씩 그룹화
        for i in range(0, len(law_names), 3):
            # 현재 페이지의 조례들
            current_laws = law_names[i:i+3]
            current_articles = all_articles[i:i+3]
            # 3개 미만이면 빈 값으로 채움
            while len(current_laws) < 3:
                current_laws.append("")
                current_articles.append([])
            # 표 생성 (1행, 3열 고정)
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            # 각 셀에 조례 내용 추가
            for idx, (law_name, articles) in enumerate(zip(current_laws, current_articles)):
                cell = table.cell(0, idx)
                paragraph = cell.paragraphs[0]
                if law_name:
                    run = paragraph.add_run(law_name + '\n')
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                    if articles:
                        paragraph.add_run('\n'.join(articles))
            # 마지막 페이지가 아니면 페이지 나누기 추가
            if i + 3 < len(law_names):
                doc.add_page_break()
                section = doc.sections[-1]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Mm(420)
                section.page_height = Mm(297)
        
        now = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"조례_검색결과_{now}.docx"
        doc.save(filename)
        # 완료 메시지
        if hasattr(self, 'result_text'):
            self.result_text.insert('1.0', f"MS워드 저장이 완료되었습니다. (파일명: {filename})\n")

    def save_selected_to_word(self):
        search_query = self.search_entry.get().strip()
        if not search_query:
            return
        law_names = []
        all_articles = []
        for org_code, metro_name in self.metropolitan_govs.items():
            params = {
                'OC': self.OC,
                'target': 'ordin',
                'type': 'XML',
                'query': search_query,
                'display': 100,
                'search': 1,
                'sort': 'ddes',
                'page': 1,
                'org': org_code
            }
            try:
                response = requests.get(self.search_url, params=params)
                root = ET.fromstring(response.text)
                for law in root.findall('.//law'):
                    ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                    ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                    기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                    if 기관명 != metro_name:
                        continue  # 본청이 아니면 건너뜀
                    search_terms = [term.lower() for term in search_query.split() if term.strip()]
                    ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                    if not all(term in ordinance_name_clean for term in search_terms):
                        continue
                    articles = self.get_ordinance_detail(ordinance_id)
                    law_names.append(ordinance_name)
                    all_articles.append(articles)
            except Exception as e:
                continue
        if law_names and all_articles:
            self.save_to_word(law_names, all_articles)
            if hasattr(self, 'result_text'):
                self.result_text.insert('1.0', "MS워드 저장이 완료되었습니다.\n")

    def upload_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            self.pdf_file_path = file_path
            self.result_text.insert('1.0', f"PDF 파일 업로드 완료: {file_path}\n")
            self.result_text.insert('1.0', "PDF 업로드가 완료되었습니다.\n")
        else:
            self.pdf_file_path = None

    def extract_pdf_text(self, pdf_path):
        text = ""
        try:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            text = f"PDF 추출 오류: {e}"
        return text

    def debug_compare_pdf_with_laws(self):
        try:
            self.result_text.insert(tk.END, '[DEBUG] compare_pdf_with_laws 버튼 클릭됨\n')
            self.compare_pdf_with_laws()
        except Exception as e:
            self.result_text.insert(tk.END, f'[DEBUG] compare_pdf_with_laws 최상위 예외: {e}\n')

    def compare_pdf_with_laws(self):
        self.result_text.insert(tk.END, '[DEBUG] compare_pdf_with_laws 함수 진입\n')
        # 본청 조례명/조문 리스트가 저장되어 있는지 확인
        law_names = getattr(self, 'last_search_law_names', [])
        all_articles = getattr(self, 'last_search_all_articles', [])
        if not law_names or not all_articles:
            self.result_text.insert(tk.END, "먼저 조례 검색을 해주세요.\n")
            return
        search_query = self.search_entry.get().strip()
        if not hasattr(self, 'pdf_file_path') or not self.pdf_file_path:
            self.result_text.insert(tk.END, "\nPDF 파일을 먼저 업로드해주세요.\n")
            return
        if not search_query:
            self.result_text.insert(tk.END, "\n검색어를 입력해주세요.\n")
            return
        self.result_text.insert(tk.END, "비교 분석 및 저장을 시작합니다. 잠시만 기다려주세요...\n")
        self.result_text.insert(tk.END, f"[DEBUG] PDF 파일 경로: {self.pdf_file_path}\n")
        pdf_text = self.extract_pdf_text(self.pdf_file_path)
        self.result_text.insert(tk.END, f"[DEBUG] PDF 텍스트 길이: {len(pdf_text)}\n")
        try:
            self.result_text.insert(tk.END, f"[DEBUG] 본청 조례 비교 분석 시작 (조례 수: {len(law_names)})\n")
            gemini_result = self.compare_with_gemini(pdf_text, law_names, all_articles)
            if not gemini_result:
                self.result_text.insert(tk.END, "\nGemini 분석 결과가 없습니다.\n")
                return
            self.result_text.insert(tk.END, f"[DEBUG] Gemini 분석 결과 수신, 워드 저장 시작\n")
            self.save_gemini_comparison_to_word(law_names, all_articles, gemini_result)
            self.result_text.insert(tk.END, f"[DEBUG] 워드 저장 완료\n")
        except Exception as e:
            self.result_text.insert(tk.END, f"\n비교 중 오류 발생: {e}\n")

    def compare_with_gemini(self, pdf_text, law_names, all_articles):
        self.result_text.insert(tk.END, f"[DEBUG] compare_with_gemini 호출\n")
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = (
            "아래는 내가 업로드한 조례 PDF의 전체 내용이야.\n"
            "---\n"
            f"{pdf_text}\n"
            "---\n"
            "그리고 아래는 타시도 조례명과 각 조문 내용이야.\n"
        )
        for name, articles in zip(law_names, all_articles):
            prompt += f"조례명: {name}\n"
            for idx, article in enumerate(articles):
                prompt += f"제{idx+1}조: {article}\n"
        prompt += (
            "---\n"
            "아래 기준에 따라 분석해줘. 반드시 한글로 답변해줘.\n"
            "1. [비교분석 요약표(조문별)]\n"
            "- 표의 컬럼: 조문(내 조례), 주요 내용, 타 시도 유사 조항, 동일 여부, 차이 및 내 조례 특징, 추천 조문\n"
            "- 반드시 내 조례(PDF로 업로드한 조례)의 조문만을 기준으로, 각 조문별로 타 시도 조례와 비교해 표로 정리(내 조례에 없는 조문은 비교하지 말 것)\n"
            "- '추천 조문' 칸에는 타 시도 조례와 비교해 무난하게 생각되는 조문 예시를 한글로 작성\n\n"
            "2. [내 조례의 차별점 요약] (별도 소제목)\n"
            "- 타 시도 조례와 비교해 독특하거나 구조적으로 다른 점, 내 조례만의 관리/운영 방식 등 요약\n\n"
            "3. [검토 시 유의사항] (별도 소제목)\n"
            "다음 원칙들을 기준으로 검토해줘:\n"
            "a) 소관사무의 원칙\n"
            "- 지방자치단체의 자치사무와 법령에 의해 위임된 단체위임사무에 대해서만 제정 가능한지\n"
            "- 국가사무가 지방자치단체의 장에게 위임된 기관위임사무인 경우 조례 제정이 적절한지\n"
            "- 사무의 성격이 전국적으로 통일적 처리를 요구하는지 여부 검토\n\n"
            "b) 법률 유보의 원칙\n"
            "- 주민의 권리 제한 또는 의무 부과에 관한 사항이나 벌칙을 정할 때 법률의 위임이 있는지\n"
            "- 법률의 위임 없이 인·허가, 등록, 신고 등을 신설하거나 정년 제한, 의무사항 신설 등이 있는지\n"
            "- 추상적인 노력 의무나 책무 부과가 사실상 강제하는 것이 되는지\n\n"
            "c) 집행기관과 의결기관 간 견제와 균형의 원리\n"
            "- 지방의회와 지방자치단체의 장의 고유권한을 침해하는지\n"
            "- 예산 편성을 강제하는 등 예산 편성권을 실체적으로 제약하는지\n"
            "- 행정기구 설치권 등 집행기관의 고유권한을 침해하는지\n\n"
            "d) 법의 일반 원칙 준수 여부\n"
            "- 비례의 원칙: 입법 목적의 정당성, 방법의 적정성, 피해의 최소성, 법익의 균형성\n"
            "- 평등의 원칙: 합리적인 이유 없는 자의적인 차별이 있는지\n"
            "- 신뢰보호의 원칙: 소급 적용 시 침해적인 경우가 있는지\n"
            "- 적법절차의 원칙: 행정절차의 적법성과 「지방자치법」의 절차 준수 여부\n"
            "- 체계정당성의 원칙: 규범 간 상호 배치나 모순이 있는지\n\n"
            "e) 포괄위임금지의 원칙\n"
            "- 조례로 위임한 사항을 규칙에 포괄적으로 재위임하는지\n"
            "- 주민에게 중요한 의미를 가지거나 제도의 핵심이 되는 사항이 조례에 직접 규정되어 있는지\n"
        )
        try:
            self.result_text.insert(tk.END, f"[DEBUG] Gemini API 호출 시작\n")
            response = model.generate_content(prompt)
            self.result_text.insert(tk.END, f"[DEBUG] Gemini API 응답 수신\n")
            return response.text
        except Exception as e:
            self.result_text.insert(tk.END, f"Gemini API 오류: {e}\n")
            return None

    def save_gemini_comparison_to_word(self, law_names, all_articles, gemini_result):
        from datetime import datetime
        from docx.shared import Mm, RGBColor, Pt
        from docx.enum.section import WD_ORIENT
        import re
        doc = Document()
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(420)
        section.page_height = Mm(297)
        doc.add_heading("비교분석", level=1)
        doc.add_paragraph("1. 비교분석 요약표(조문별)")

        # 1. 비교분석 요약표(조문별) 추출 및 표로 변환
        table_pattern = re.compile(r'(\|.+\|\n)+')
        table_match = table_pattern.search(gemini_result)
        table_text = ''
        if table_match:
            table_text = table_match.group()
            rows = [row.strip() for row in table_text.strip().split('\n') if row.strip()]
            # 마크다운 구분선(---) 제거
            rows = [row for row in rows if not set(row.replace('|','').strip()) <= set('-')]
            table_data = [ [cell.strip().replace('**','') for cell in row.split('|')[1:-1]] for row in rows ]
            # 표 생성: 추천 조문 칸 추가(마지막 열)
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            # 각 셀에 내용 추가 및 칸 너비 조정
            for i, cell in enumerate(table.rows[0].cells):
                cell.text = table_data[0][i]
                # '동일 여부' 칸은 좁게
                if '동일 여부' in cell.text:
                    cell.width = Mm(20)
                # '추천 조문' 칸은 넓게
                if '추천 조문' in cell.text:
                    cell.width = Mm(80)
            for row in table_data[1:]:
                cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    cells[i].text = row[i]
            doc.add_paragraph('')
            # 표 이후 텍스트만 남기기 위해 결과에서 표 부분 제거
            gemini_result = gemini_result.replace(table_text, '')

        # 2. 차별점 요약, 3. 검토시 유의사항 등 나머지 텍스트(마크다운 기호 제거)
        clean_text = re.sub(r'[#*`>\-]+', '', gemini_result)
        for line in clean_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        # 4. 최종 검토: 상위법 위반 여부 검토 필요성이 언급된 줄에서만 상위법령 후보 추출
        upper_law_candidates = set()
        law_pattern = re.compile(r'([\w가-힣]+(법|시행령|시행규칙))')
        # 상위법 위반 여부 검토 필요성 키워드
        law_check_keywords = [
            '위반', '위배', '충돌', '저촉', '준수', '적합', '불일치',
            '상위법', '상위 법령', '상위법령', '법령과의 관계', '법령과의 충돌', '법령과의 위배'
        ]
        # 비교분석 요약표(조문별)에서 추출
        for row in table_text.split('\n'):
            if any(keyword in row for keyword in law_check_keywords):
                for m in law_pattern.finditer(row):
                    upper_law_candidates.add(m.group(1))
        # 차별점 요약, 검토시 유의사항 등 나머지 텍스트에서 추출
        for line in clean_text.split('\n'):
            if any(keyword in line for keyword in law_check_keywords):
                for m in law_pattern.finditer(line):
                    upper_law_candidates.add(m.group(1))
        pdf_text = self.extract_pdf_text(self.pdf_file_path)
        # 1. 불용어 리스트(실존하지 않는 법령명)
        invalid_law_names = {'자치입법', '조례', '규칙', '지침', '내규', '예규', '훈령', '적법', '입법', '상위법', '위법', '합법', '불법', '방법', '헌법상', '헌법적', '법적', '법률적', '법령상', '법률상', '법률', '법령', '법', '규정', '조항', '조문', '규범', '원칙', '기준', '사항', '내용', '관련법', '관련 법', '관련법령', '관련 법령'}
        def is_valid_law_name(name):
            # 대소문자, 공백 모두 제거 후 비교
            name_clean = name.strip().replace(' ', '').lower()
            for invalid in invalid_law_names:
                if name_clean == invalid.replace(' ', '').lower():
                    print(f"[DEBUG] 불용어 '{name}'(clean: '{name_clean}')는 query로 검색하지 않음.")
                    return False
            # 숫자+법(예: 1법, 2법 등)도 제외
            if name_clean and name_clean[0].isdigit():
                print(f"[DEBUG] 숫자+법 '{name}'(clean: '{name_clean}')는 query로 검색하지 않음.")
                return False
            return True
        for upper_law_name in upper_law_candidates:
            if not is_valid_law_name(upper_law_name):
                continue  # 실존하지 않는 법령명 또는 불용어는 건너뜀
            # 1. lawSearch로 현행 법령ID 및 법령명한글 얻기
            search_url = 'http://www.law.go.kr/DRF/lawSearch.do'
            search_params = {
                'OC': self.OC,
                'target': 'law',
                'type': 'XML',
                'query': upper_law_name
            }
            print(f"[DEBUG] lawSearch 요청 URL: {search_url}")
            print(f"[DEBUG] lawSearch 요청 파라미터: {search_params}")
            search_resp = requests.get(search_url, params=search_params)
            print(f"[DEBUG] lawSearch 응답코드: {search_resp.status_code}")
            print(f"[DEBUG] lawSearch 응답 본문(앞 1000자): {search_resp.text[:1000]}")
            search_root = ET.fromstring(search_resp.text)
            law_id = None
            law_name_kor = None
            # 여러 개가 검색될 경우 반드시 '현행'인 것만 골라서 사용
            for law in search_root.findall('.//law'):
                if law.find('현행연혁코드') is not None and law.find('현행연혁코드').text == '현행':
                    law_id = law.find('법령ID').text if law.find('법령ID') is not None else None
                    law_name_kor = law.find('법령명한글').text if law.find('법령명한글') is not None else None
                    break
            if not law_id or not law_name_kor:
                print(f"[DEBUG] 현행 법령ID 또는 법령명한글을 찾을 수 없음: {upper_law_name}")
                continue  # 실존하지 않는 법령명은 건너뜀
            # 2. lawService로 본문 요청 (조문내용만 추출)
            detail_url = 'http://www.law.go.kr/DRF/lawService.do'
            detail_params = {
                'OC': self.OC,
                'target': 'law',
                'type': 'XML',
                'ID': law_id
            }
            print(f"[DEBUG] lawService 요청 URL: {detail_url}")
            print(f"[DEBUG] lawService 요청 파라미터: {detail_params}")
            detail_resp = requests.get(detail_url, params=detail_params)
            print(f"[DEBUG] lawService 응답코드: {detail_resp.status_code}")
            print(f"[DEBUG] lawService 응답 본문(앞 1000자): {detail_resp.text[:1000]}")
            detail_root = ET.fromstring(detail_resp.text)
            upper_law_text = ''
            for article in detail_root.findall('.//조'):
                content = None
                # 다양한 태그명을 시도하여 본문 추출
                for tag in ['조문내용', '조내용', '내용']:
                    node = article.find(tag)
                    if node is not None and node.text:
                        content = node.text
                        break
                if content:
                    content = content.replace('<![CDATA[', '').replace(']]>', '').replace('<p>', '').replace('</p>', '\n').replace('<br/>', '\n').replace('<br>', '\n').replace('&nbsp;', ' ').strip()
                    upper_law_text += content + '\n'
            # Gemini에 위배 여부 추가 분석 요청
            doc.add_heading('상위법령 위반 여부 검토', level=2)
            doc.add_paragraph(f'상위 법령명: {upper_law_name}')
            doc.add_paragraph('상위 법령 본문 일부:')
            doc.add_paragraph(upper_law_text[:2000])
            # Gemini 추가 분석
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f'아래는 상위 법령({upper_law_name})의 주요 내용과 내가 업로드한 조례의 전체 내용이야.\n'
                '---상위 법령---\n'
                f'{upper_law_text[:2000]}\n'
                '---내 조례---\n'
                f'{pdf_text[:2000]}\n'
                '---\n'
                '다음 기준에 따라 분석해줘. 반드시 한글로 답변해줘.\n'
                '1. [법령우위의 원칙 위반 여부]\n'
                '- 조례가 상위 법령의 내용과 직접적으로 충돌하거나 위배되는지\n'
                '- 상위 법령의 취지나 목적을 해치는지\n'
                '- 상위 법령이 금지하는 행위를 허용하거나, 의무화하는 행위를 면제하는지\n\n'
                '2. [법률 유보의 원칙 위반 여부]\n'
                '- 주민의 권리를 제한하거나 의무를 부과하는 내용이 있는지\n'
                '- 상위 법령에서 위임받지 않은 권한을 행사하는지\n'
                '- 상위 법령의 위임 범위를 초과하는지\n\n'
                '3. [실무적 검토 포인트]\n'
                '- 조례의 집행 과정에서 발생할 수 있는 문제점\n'
                '- 상위 법령과의 관계에서 주의해야 할 사항\n'
                '- 개선이 필요한 부분과 그 방향성\n'
            )
            try:
                response = model.generate_content(prompt)
                clean_gemini = re.sub(r'[\*#`>\-]+', '', response.text)
                for line in clean_gemini.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            except Exception as e:
                doc.add_paragraph(f'Gemini API 오류: {e}')

        now = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"조례_Gemini_비교_{now}.docx"
        doc.save(filename)
        # 완료 메시지
        if hasattr(self, 'result_text'):
            self.result_text.insert('1.0', f"비교 분석 및 MS워드 저장이 완료되었습니다. (파일명: {filename})\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = LawSearchApp(root)
    root.mainloop() 