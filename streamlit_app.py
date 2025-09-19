import streamlit as st
import requests
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
import PyPDF2
import google.generativeai as genai
import openai
import os
import tempfile
import re
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import numpy as np
import pickle
import hashlib
from typing import Dict, List
from sklearn.metrics.pairwise import cosine_similarity

# 페이지 설정
st.set_page_config(
    page_title="광역지자체 조례 검색, 비교, 분석",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 사용자 정의 CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #4f46e5, #7c3aed);
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        color: white;
        margin-bottom: 2rem;
    }
    .step-card {
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .result-card {
        background: #ffffff;
        border: 1px solid #d1d5db;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .law-title {
        color: #dc2626;
        font-weight: bold;
    }

    /* 탭 글자 크기 키우기 */
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
        font-size: 18px !important;
        font-weight: 600 !important;
        font-size: 1.1em;
        margin-bottom: 0.5rem;
    }
    .metro-name {
        color: #1e40af;
        font-weight: 600;
        margin-bottom: 0.3rem;
    }
    .stButton > button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

# API 설정
OC = "climsneys85"
search_url = "http://www.law.go.kr/DRF/lawSearch.do"
detail_url = "http://www.law.go.kr/DRF/lawService.do"
precedent_search_url = "http://www.law.go.kr/DRF/lawSearch.do"  # 판례 검색 API

# 광역지자체 코드 및 이름
metropolitan_govs = {
    '6110000': '서울특별시',
    '6260000': '부산광역시',
    '6270000': '대구광역시',
    '6280000': '인천광역시',
    '6290000': '광주광역시',
    '6300000': '대전광역시',
    '5690000': '세종특별자치시',
    '6310000': '울산광역시',
    '6410000': '경기도',
    '6530000': '강원특별자치도',
    '6430000': '충청북도',
    '6440000': '충청남도',
    '6540000': '전북특별자치도',
    '6460000': '전라남도',
    '6470000': '경상북도',
    '6480000': '경상남도',
    '6500000': '제주특별자치도'
}

# 세션 상태 초기화
if 'search_results' not in st.session_state:
    st.session_state.search_results = []
if 'uploaded_pdf' not in st.session_state:
    st.session_state.uploaded_pdf = None
if 'search_query' not in st.session_state:
    st.session_state.search_query = ""
if 'word_doc_ready' not in st.session_state:
    st.session_state.word_doc_ready = False
if 'word_doc_data' not in st.session_state:
    st.session_state.word_doc_data = None
if 'selected_ordinances' not in st.session_state:
    st.session_state.selected_ordinances = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None

def get_ordinance_detail(ordinance_id):
    """조례 상세 내용 가져오기"""
    params = {
        'OC': OC,
        'target': 'ordin',
        'ID': ordinance_id,
        'type': 'XML'
    }
    try:
        response = requests.get(detail_url, params=params, timeout=60)
        root = ET.fromstring(response.text)
        articles = []
        for article in root.findall('.//조'):
            content = article.find('조내용').text if article.find('조내용') is not None else ""
            if content:
                content = content.replace('<![CDATA[', '').replace(']]>', '')
                content = content.replace('<p>', '').replace('</p>', '\n')
                content = content.replace('<br/>', '\n')
                content = content.replace('<br>', '\n')
                content = content.replace('&nbsp;', ' ')
                content = content.strip()
            if content:
                articles.append(content)
        return articles
    except Exception:
        return []

def search_ordinances(query):
    """조례 검색 함수"""
    results = []
    total_count = 0
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    total_metros = len(metropolitan_govs)
    
    for idx, (org_code, metro_name) in enumerate(metropolitan_govs.items()):
        status_text.text(f"검색 중... {metro_name} ({idx + 1}/{total_metros})")
        progress_bar.progress((idx + 1) / total_metros)
        
        try:
            params = {
                'OC': OC,
                'target': 'ordin',
                'type': 'XML',
                'query': query,
                'display': 100,
                'search': 1,
                'sort': 'ddes',
                'page': 1,
                'org': org_code
            }
            
            response = requests.get(search_url, params=params, timeout=60)
            response.raise_for_status()
            
            root = ET.fromstring(response.text)
            
            for law in root.findall('.//law'):
                ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                
                if 기관명 != metro_name:
                    continue
                
                # 검색어 매칭 로직
                search_terms = [term.lower() for term in query.split() if term.strip()]
                ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                if not all(term in ordinance_name_clean for term in search_terms):
                    continue
                
                total_count += 1
                articles = get_ordinance_detail(ordinance_id)
                
                results.append({
                    'name': ordinance_name,
                    'content': articles,
                    'metro': metro_name
                })
                
        except Exception as e:
            st.warning(f"검색 중 오류 발생 ({metro_name}): {str(e)}")
            continue
    
    progress_bar.empty()
    status_text.empty()
    
    return results, total_count

def create_word_document(query, results):
    """Word 문서 생성 함수"""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)

    # 제목 추가
    title = doc.add_heading('조례 검색 결과', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'검색어: {query}')
    doc.add_paragraph(f'총 {len(results)}건의 조례가 검색되었습니다.\n')

    # 조례를 3개씩 그룹화하여 3단 비교표 형태로 생성
    for i in range(0, len(results), 3):
        current_laws = results[i:i+3]
        while len(current_laws) < 3:
            current_laws.append({'name': '', 'content': [], 'metro': ''})

        # 표 생성 (1행, 3열 고정)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = True

        # 각 셀에 조례 내용 추가
        for idx, law in enumerate(current_laws):
            cell = table.cell(0, idx)
            paragraph = cell.paragraphs[0]
            
            if law['name']:
                # 조례명 추가 (지자체명 + 조례명)
                run = paragraph.add_run(f"{law['metro']}\n{law['name']}\n\n")
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                
                # 조문 내용 추가
                if law['content']:
                    content_text = '\n\n'.join(law['content'])
                    paragraph.add_run(content_text)
                else:
                    paragraph.add_run('(조문 없음)')

        # 마지막 페이지가 아니면 페이지 나누기 추가
        if i + 3 < len(results):
            doc.add_page_break()

    return doc

def extract_pdf_text(pdf_file):
    """PDF 텍스트 추출 함수"""
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
        return text
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류 발생: {str(e)}")
        return None

def extract_superior_laws(pdf_text):
    """조례안에서 상위법령 추출 함수 - GUI 검증된 로직 적용"""
    import re

    # 상위법 후보 추출을 위한 키워드 (조례안에서 상위법령 언급하는 모든 맥락 포함)
    law_check_keywords = [
        '위반', '위배', '충돌', '저촉', '준수', '적합', '불일치',
        '상위법', '상위 법령', '상위법령', '법령과의 관계', '법령과의 충돌', '법령과의 위배',
        '관계법령', '근거법령', '법적근거', '참고사항', '관련법령', '소관법령',
        '법령', '법률', '시행령', '시행규칙', '규정', '개정', '제정', '법'  # 일반적인 법령 언급
    ]

    # 법령명 패턴 (시행령/시행규칙 추출 개선)
    law_pattern = re.compile(r'([가-힣\w\s]*(?:법|시행령|시행규칙))\s*(?:[」]|$|[.,;:\s])', re.MULTILINE)

    # 상위법 후보 추출
    upper_law_candidates = set()

    # 1. 상위법 관련 맥락이 있는 줄에서 법령명 추출
    for line in pdf_text.split('\n'):
        if any(keyword in line for keyword in law_check_keywords):
            for match in law_pattern.finditer(line):
                law_name = match.group(1).strip()
                if law_name:
                    upper_law_candidates.add(law_name)

    # 2. 추가 패턴: 「법령명」 형식으로 따옴표 안에 있는 법령명 추출
    quote_pattern = re.compile(r'[「『]([^」』]*(?:법|시행령|시행규칙))[」』]')
    for match in quote_pattern.finditer(pdf_text):
        law_name = match.group(1).strip()
        if law_name:
            upper_law_candidates.add(law_name)

    # 3. 추가 패턴: "○○법령:" 또는 "관계법령:" 뒤에 오는 법령명
    relation_pattern = re.compile(r'(?:관계법령|근거법령|법적근거|소관법령|관련법령)\s*[:：]\s*[「『]?([^」』\n]*(?:법|시행령|시행규칙))[」』]?')
    for match in relation_pattern.finditer(pdf_text):
        law_name = match.group(1).strip()
        if law_name:
            upper_law_candidates.add(law_name)

    # 불용어 리스트 (실존하지 않는 법령명)
    invalid_law_names = {
        '자치입법', '조례', '규칙', '지침', '내규', '예규', '훈령', '적법',
        '입법', '상위법', '위법', '합법', '불법', '방법', '헌법상', '헌법적',
        '법적', '법률적', '법령상', '법률상', '법률', '법령', '법', '규정',
        '조항', '조문', '규범', '원칙', '기준', '사항', '내용', '관련법',
        '관련 법', '관련법령', '관련 법령'
    }

    def is_valid_law_name(name):
        """유효한 법령명인지 검증"""
        # 대소문자, 공백 모두 제거 후 비교
        name_clean = name.strip().replace(' ', '').lower()

        # 불용어 체크
        for invalid in invalid_law_names:
            if name_clean == invalid.replace(' ', '').lower():
                return False

        # 숫자+법(예: 1법, 2법 등)도 제외
        if name_clean and name_clean[0].isdigit():
            return False

        # 너무 짧은 이름 제외
        if len(name_clean) < 3:
            return False

        return True

    # 유효한 법령명만 필터링
    valid_laws = []
    for law_name in upper_law_candidates:
        if is_valid_law_name(law_name):
            valid_laws.append(law_name)

    # 🆕 시행령/시행규칙 자동 유추 추가
    additional_laws = []
    for law in valid_laws:
        if law.endswith('법') and '시행' not in law:
            # 해당 법률의 시행령과 시행규칙을 자동으로 추가
            base_name = law

            # 시행령 추가 (일반적인 패턴)
            potential_decree = f"{base_name} 시행령"
            if potential_decree not in valid_laws:
                additional_laws.append(potential_decree)

            # 시행규칙 추가 (일반적인 패턴)
            potential_rule = f"{base_name} 시행규칙"
            if potential_rule not in valid_laws:
                additional_laws.append(potential_rule)

    # 추가된 법령들을 포함
    if additional_laws:
        import streamlit as st
        st.info(f"🔄 자동 추가된 하위 법령: {len(additional_laws)}개")
        with st.expander("📋 자동 추가된 법령", expanded=False):
            for law in additional_laws:
                st.markdown(f"- {law}")
        valid_laws.extend(additional_laws)

    # 중복 제거 및 정렬
    unique_laws = list(set(valid_laws))
    unique_laws.sort()

    return unique_laws[:20]  # 최대 20개 반환

def get_superior_law_content_xml(law_name):
    """XML API를 통해 상위법령 내용 가져오기 (성공적인 로직 적용)"""
    try:
        import xml.etree.ElementTree as ET
        import re


        # 검색어 최적화: 띄어쓰기와 특수문자 정리
        search_query = law_name.strip()

        # 1단계: 법령 검색 (더 많은 결과 반환)
        search_params = {
            'OC': OC,
            'target': 'law',
            'type': 'XML',
            'query': search_query,
            'display': 10  # 더 많은 결과 검색
        }
        
        search_response = requests.get(search_url, params=search_params, timeout=30)
        if search_response.status_code != 200:
            st.error(f"[DEBUG] 검색 실패: HTTP {search_response.status_code}")
            return get_superior_law_content_xml_fallback(law_name)
        
        search_root = ET.fromstring(search_response.text)
        
        # 현행 법령 찾기 - 더 유연한 검색
        current_laws = []
        for law in search_root.findall('.//law'):
            status = law.find('현행연혁코드')
            if status is not None and status.text == '현행':
                law_id_elem = law.find('법령ID')
                law_name_elem = law.find('법령명한글')
                if law_id_elem is not None and law_name_elem is not None:
                    current_laws.append({
                        'id': law_id_elem.text,
                        'name': law_name_elem.text
                    })
        
        if not current_laws:
            st.warning(f"[DEBUG] {law_name}의 현행 법령을 찾을 수 없음")
            return get_superior_law_content_xml_fallback(law_name)
        
        # 가장 관련성 높은 법령 선택 (개선된 매칭 알고리즘)
        best_law = None
        best_score = -1

        for law_info in current_laws:
            found_name = law_info['name']
            score = 0

            # 1. 정확한 매칭 우선
            if found_name == law_name:
                score += 1000

            # 2. 부분 매칭 점수 (양방향)
            if law_name in found_name:
                score += 500
            if found_name in law_name:
                score += 300

            # 3. 핵심 키워드 매칭 (개선된 로직)
            law_lower = law_name.lower().replace(' ', '')
            found_lower = found_name.lower().replace(' ', '')

            # 여객자동차 운수사업법 관련 특별 점수
            if '여객자동차' in law_lower and '운수사업' in law_lower:
                if '여객자동차' in found_lower and '운수사업' in found_lower:
                    score += 400  # 여객자동차 운수사업법 관련 높은 점수
                    if '시행규칙' in law_lower and '시행규칙' in found_lower:
                        score += 200  # 시행규칙 매칭 추가 점수

            # 도로교통법 관련
            if '도로' in law_lower and '교통' in law_lower:
                if '도로교통' in found_lower and '특별회계' not in found_lower:
                    score += 300
                elif '교통시설' in found_lower:
                    score -= 100

            # 4. 법령 유형 매칭 점수 (요청된 유형과 일치하는지)
            requested_type = ''
            if '시행규칙' in law_lower:
                requested_type = '시행규칙'
            elif '시행령' in law_lower:
                requested_type = '시행령'
            elif '법' in law_lower and '시행' not in law_lower:
                requested_type = '법'

            if requested_type:
                if requested_type in found_lower:
                    score += 300  # 요청된 법령 유형과 일치하면 높은 점수
                elif requested_type == '법' and found_lower.endswith('법') and '시행' not in found_lower:
                    score += 300
            else:
                # 기본 우선순위 (법률 > 시행령 > 시행규칙)
                if found_lower.endswith('법') and not ('시행령' in found_lower or '시행규칙' in found_lower):
                    score += 100
                elif '시행령' in found_lower:
                    score += 50
                elif '시행규칙' in found_lower:
                    score += 25

            # 5. 길이 페널티 완화 (너무 긴 법령명은 약간 감점)
            if len(found_name) > 30:
                score -= 30
                
            
            if score > best_score:
                best_score = score
                best_law = law_info
        
        if best_law:
            law_id = best_law['id']
            exact_law_name = best_law['name']
        else:
            # 폴백: 첫 번째 법령
            law_id = current_laws[0]['id']
            exact_law_name = current_laws[0]['name']
        
        if not law_id:
            st.warning(f"[DEBUG] {law_name}의 현행 법령을 찾을 수 없음")
            return get_superior_law_content_xml_fallback(law_name)
        
        # 2단계: 상세 정보 가져오기
        detail_params = {
            'OC': OC,
            'target': 'law',
            'type': 'XML',
            'ID': law_id
        }
        
        detail_response = requests.get(detail_url, params=detail_params, timeout=30)
        if detail_response.status_code != 200:
            st.error(f"[DEBUG] 상세 조회 실패: HTTP {detail_response.status_code}")
            return get_superior_law_content_xml_fallback(law_name)
        
        detail_root = ET.fromstring(detail_response.text)
        
        # 3단계: 성공적인 추출 로직 적용 - 연결된 본문으로 처리
        upper_law_text = ""
        jo_count = 0
        hang_count = 0 
        ho_count = 0
        
        for node in detail_root.iter():
            if node.tag == '조문내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += content + '\n'
                jo_count += 1
            elif node.tag == '항내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += '    ' + content + '\n'
                hang_count += 1
            elif node.tag == '호내용' and node.text and node.text.strip():
                content = re.sub(r'<[^>]+>', '', node.text)
                content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                upper_law_text += '        ' + content + '\n'
                ho_count += 1
        
        
        if upper_law_text.strip():
            # 스마트 필터링: 조례 관련 키워드가 포함된 부분 우선 추출
            def smart_filter_content(content, max_length=50000):
                """조례와 관련성 높은 부분을 우선 추출"""
                lines = content.split('\n')
                
                # 조례 관련 키워드 (도로교통법 관련)
                priority_keywords = [
                    '시장', '군수', '구청장', '지방자치단체', '조례', '시도', '시군구',
                    '위임', '위탁', '권한', '사무', '신고', '허가', '승인', '지정',
                    '주차', '정차', '금지', '제한', '구역', '시설', '설치', '관리'
                ]
                
                # 우선순위별로 라인 분류
                high_priority = []
                medium_priority = []
                low_priority = []
                
                for line in lines:
                    line_lower = line.lower()
                    priority_count = sum(1 for keyword in priority_keywords if keyword in line_lower)
                    
                    if priority_count >= 2:
                        high_priority.append(line)
                    elif priority_count >= 1:
                        medium_priority.append(line)
                    else:
                        low_priority.append(line)
                
                # 우선순위별로 결합
                filtered_content = []
                current_length = 0
                
                # 1단계: 높은 우선순위
                for line in high_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                # 2단계: 중간 우선순위
                for line in medium_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                # 3단계: 낮은 우선순위 (공간이 남으면)
                for line in low_priority:
                    if current_length + len(line) < max_length:
                        filtered_content.append(line)
                        current_length += len(line)
                    else:
                        break
                
                result = '\n'.join(filtered_content)
                if len(content) > len(result):
                    result += "\n\n[... 조례 관련성이 높은 부분을 우선 표시하였습니다 ...]"
                
                return result
            
            # 스마트 필터링 적용 (Gemini 2.0 flash exp는 더 큰 컨텍스트 지원)
            max_length = 80000
            if len(upper_law_text) > max_length:
                truncated_text = smart_filter_content(upper_law_text, max_length)
                st.warning(f"[DEBUG] 본문이 너무 길어 조례 관련 부분을 우선하여 {len(truncated_text):,}자로 축약했습니다 (원본: {len(upper_law_text):,}자)")
            else:
                truncated_text = upper_law_text.strip()
            
            # 모든 조문을 하나의 연결된 본문으로 처리
            result = {
                'law_name': exact_law_name,
                'law_id': law_id,
                'content': truncated_text
            }
            
            return result
        else:
            st.warning("[DEBUG] 조문 내용이 비어있음")
            return get_superior_law_content_xml_fallback(law_name)
        
    except Exception as e:
        st.error(f"[DEBUG] 예외 발생: {str(e)}")
        return get_superior_law_content_xml_fallback(law_name)

def get_superior_law_content_xml_fallback(law_name):
    """XML 방식 폴백 (간소화 버전)"""
    try:
        st.write(f"[DEBUG XML] XML 폴백 모드 시작: {law_name}")
        
        search_params = {
            'OC': OC, 
            'target': 'law',
            'type': 'XML',
            'query': law_name,
            'display': 5,
            'search': 1
        }
        
        st.write(f"[DEBUG XML] XML 검색 파라미터: {search_params}")
        
        search_response = requests.get(search_url, params=search_params, timeout=30)
        st.write(f"[DEBUG XML] XML 검색 응답 상태: {search_response.status_code}")
        st.write(f"[DEBUG XML] 응답 내용 (처음 1000자): {search_response.text[:1000]}")
        
        if search_response.status_code != 200:
            st.error(f"[DEBUG XML] XML API 호출 실패: HTTP {search_response.status_code}")
            return None
            
        if not search_response.text.strip():
            st.error("[DEBUG XML] XML 응답이 비어있습니다")
            return None
            
        try:
            search_root = ET.fromstring(search_response.text)
        except ET.ParseError as xml_err:
            st.error(f"[DEBUG XML] XML 파싱 실패: {xml_err}")
            st.write(f"[DEBUG XML] 원본 응답: {search_response.text}")
            return None
        st.write(f"[DEBUG XML] XML 파싱 완료")
        
        law_id = None
        exact_law_name = None
        
        for law in search_root.findall('.//law'):
            found_name = law.find('법령명').text if law.find('법령명') is not None else ""
            found_id = law.find('법령ID').text if law.find('법령ID') is not None else None
            
            if found_name == law_name or (law_name in found_name):
                law_id = found_id
                exact_law_name = found_name
                break
        
        if not law_id:
            return None
        
        detail_params = {
            'OC': OC,
            'target': 'law', 
            'ID': law_id,
            'type': 'XML'
        }
        
        detail_response = requests.get(detail_url, params=detail_params, timeout=30)
        detail_root = ET.fromstring(detail_response.text)
        
        articles = []
        for article in detail_root.findall('.//조'):
            article_num = article.find('조문번호').text if article.find('조문번호') is not None else ""
            article_title = article.find('조문제목').text if article.find('조문제목') is not None else ""
            article_content = article.find('조문내용').text if article.find('조문내용') is not None else ""
            
            if article_content:
                article_content = article_content.replace('<![CDATA[', '').replace(']]>', '')
                article_content = article_content.replace('<p>', '').replace('</p>', '\n')
                article_content = article_content.replace('<br/>', '\n').replace('<br>', '\n')
                article_content = article_content.replace('&nbsp;', ' ')
                article_content = article_content.strip()
                
                if article_content:
                    articles.append({
                        'number': article_num,
                        'title': article_title,
                        'content': article_content
                    })
        
        return {
            'law_name': exact_law_name,
            'law_id': law_id,
            'articles': articles
        }
        
    except Exception as e:
        st.error(f"[DEBUG XML] XML 폴백 예외 발생: {str(e)}")
        return None

# 기존 함수를 새 XML 방식으로 교체
def get_superior_law_content(law_name):
    """상위법령 내용 가져오기 (XML 방식)"""
    return get_superior_law_content_xml(law_name)

def normalize_law_name(law_name):
    """법령명을 정규화하여 중복 제거"""
    import re

    # 1. 기본 정리: 앞뒤 공백 제거
    normalized = law_name.strip()

    # 2. 과도한 띄어쓰기 제거 (2개 이상의 공백을 1개로)
    normalized = re.sub(r'\s+', ' ', normalized)

    # 3. 특정 패턴 정규화
    # "관광진흥 법" -> "관광진흥법"
    normalized = re.sub(r'(\w+)\s+(법|령|규칙)$', r'\1\2', normalized)

    # 4. 폐광지역개발지원 관련 법령 정규화
    if '폐광지' in normalized or '역개발' in normalized:
        if '특별법' in normalized:
            normalized = "폐광지역개발지원에관한특별법"

    # 5. 너무 짧은 법령명 제거 (3글자 이하)
    if len(normalized) <= 3:
        return None

    # 6. 명확히 잘못된 추출 제거
    invalid_patterns = [
        r'^한특별법$',  # "한특별법"
        r'^\w{1,2}특별법$',  # 너무 짧은 특별법
    ]

    for pattern in invalid_patterns:
        if re.match(pattern, normalized):
            return None

    return normalized

def group_laws_by_hierarchy(superior_laws):
    """법령을 계층별로 그룹화하는 함수 (정규화 적용)"""
    law_groups = {}

    # 1단계: 법령명 정규화 및 중복 제거
    normalized_laws = set()
    for law_name in superior_laws:
        normalized = normalize_law_name(law_name)
        if normalized:  # None이 아닌 경우만 추가
            normalized_laws.add(normalized)

    if len(superior_laws) != len(normalized_laws):
        import streamlit as st
        st.info(f"🔧 법령명 정규화: {len(superior_laws)}개 → {len(normalized_laws)}개로 중복 제거")

        # 제거된 중복 법령 표시
        removed_laws = []
        for original in superior_laws:
            normalized = normalize_law_name(original)
            if not normalized or (normalized != original and normalized in normalized_laws):
                removed_laws.append(original)

        if removed_laws:
            with st.expander("🗑️ 제거된 중복/잘못된 법령명", expanded=False):
                for removed in removed_laws:
                    st.markdown(f"- {removed}")

    # 정규화 과정 로깅
    for original in superior_laws:
        normalized = normalize_law_name(original)
        if normalized != original:
            pass  # 정규화된 경우 처리 (디버깅 코드 제거됨)

    # 2단계: 정규화된 법령명으로 그룹화
    for law_name in normalized_laws:
        # 기본 법령명 추출 (시행령, 시행규칙 제거)
        base_name = law_name
        law_type = 'law'  # 기본값: 법률
        
        if '시행규칙' in law_name:
            base_name = law_name.replace(' 시행규칙', '').replace('시행규칙', '')
            law_type = 'rule'
        elif '시행령' in law_name:
            base_name = law_name.replace(' 시행령', '').replace('시행령', '')
            law_type = 'decree'
        elif law_name.endswith('령') and not law_name.endswith('법령'):
            law_type = 'decree'
        elif law_name.endswith('규칙'):
            law_type = 'rule'
            
        # 그룹에 추가
        if base_name not in law_groups:
            law_groups[base_name] = {'law': None, 'decree': None, 'rule': None}
        
        law_groups[base_name][law_type] = law_name
    
    return law_groups

def get_all_superior_laws_content(superior_laws):
    """모든 상위법령 내용을 가져오는 함수 - 계층별 그룹화"""
    superior_laws_content = []
    
    if not superior_laws:
        return superior_laws_content
    
    # 1단계: 법령을 계층별로 그룹화
    law_groups = group_laws_by_hierarchy(superior_laws)
    
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_laws = sum(1 for laws in law_groups.values() for law in laws.values() if law is not None)
    current_idx = 0
    
    # 2단계: 각 그룹의 모든 계층 수집
    for base_name, laws in law_groups.items():
        group_content = {
            'base_name': base_name,
            'laws': {},
            'combined_articles': []
        }
        
        # 법률 → 시행령 → 시행규칙 순서로 수집
        for law_type in ['law', 'decree', 'rule']:
            law_name = laws[law_type]
            if law_name:
                current_idx += 1
                status_text.text(f"상위법령 조회 중... {law_name} ({current_idx}/{total_laws})")
                progress_bar.progress(current_idx / total_laws)
                
                law_content = get_superior_law_content(law_name)
                if law_content:
                    group_content['laws'][law_type] = law_content
                    # 새로운 데이터 구조 처리: content가 있으면 사용, articles가 있으면 변환
                    if 'content' in law_content:
                        # 연결된 본문이 있으면 그대로 저장
                        if 'combined_content' not in group_content:
                            group_content['combined_content'] = ""
                        group_content['combined_content'] += law_content['content'] + '\n'
                    elif 'articles' in law_content:
                        # 기존 articles 구조가 있으면 변환
                        group_content['combined_articles'].extend(law_content['articles'])
        
        if group_content['laws']:  # 하나 이상의 법령이 수집된 경우만 추가
            superior_laws_content.append(group_content)
    
    progress_bar.empty()
    status_text.empty()
    
    # 텍스트 길이 제한 (8만자) 및 관련성 필터링
    max_chars = 80000
    total_chars = 0
    
    # 각 법령 그룹의 텍스트 길이 계산
    for group in superior_laws_content:
        group_chars = 0
        
        # combined_content가 있는 경우
        if 'combined_content' in group and group['combined_content']:
            group_chars += len(group['combined_content'])
        
        # combined_articles가 있는 경우
        if 'combined_articles' in group and group['combined_articles']:
            for article in group['combined_articles']:
                group_chars += len(article.get('content', ''))
        
        # laws 구조가 있는 경우
        if 'laws' in group and group['laws']:
            for law_type, law_info in group['laws'].items():
                if law_info and 'articles' in law_info:
                    for article in law_info['articles']:
                        group_chars += len(article.get('content', ''))
        
        group['text_length'] = group_chars
        total_chars += group_chars
        
    
    # 8만자를 초과하는 경우 경고만 표시하고 모든 내용 유지 (필터링 비활성화)
    if total_chars > max_chars:
        st.warning(f"⚠️ 법령 내용이 {total_chars:,}자로 8만자를 초과합니다. 하지만 모든 내용을 유지합니다.")
        st.info("💡 Gemini가 긴 텍스트를 처리할 수 있으므로 필터링하지 않고 전체 내용을 전달합니다.")
    
    st.success(f"✅ 총 {len(superior_laws_content)}개 법령 그룹, {total_chars:,}자를 Gemini에게 전달합니다.")
    
    return superior_laws_content

def chunk_text(text, chunk_size=1000, overlap=200):
    """텍스트를 청크로 분할하는 함수"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = text[start:end]
        
        # 문장 단위로 끝나도록 조정
        if end < text_length:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            last_break = max(last_period, last_newline)
            if last_break > start + chunk_size * 0.7:  # 너무 짧지 않으면 조정
                end = start + last_break + 1
                chunk = text[start:end]
        
        if chunk.strip():
            chunks.append({
                'text': chunk.strip(),
                'start': start,
                'end': end
            })
        
        start = end - overlap
    
    return chunks

def get_gemini_embedding(text, api_key):
    """Gemini를 사용하여 텍스트 임베딩 생성"""
    try:
        genai.configure(api_key=api_key)
        result = genai.embed_content(
            model="models/embedding-001",
            content=text,
            task_type="retrieval_document"
        )
        return result['embedding']
    except Exception as e:
        st.error(f"임베딩 생성 오류: {str(e)}")
        return None

def create_vector_store(pdf_path, api_key):
    """자치법규 가이드 PDF를 벡터스토어로 변환 (pickle 방식)"""
    try:
        vector_store_path = "jachi_guide_2022_vectorstore.pkl"
        
        # 기존 벡터스토어가 있으면 로드
        if os.path.exists(vector_store_path):
            with open(vector_store_path, 'rb') as f:
                vector_store = pickle.load(f)
            st.info("✅ 기존 벡터스토어를 로드했습니다.")
            return vector_store
        
        # PDF 텍스트 추출
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ''
            for page in reader.pages:
                full_text += page.extract_text() + '\n'
        
        # 텍스트 청킹
        chunks = chunk_text(full_text)
        st.info(f"📄 {len(chunks)}개의 텍스트 청크로 분할했습니다.")
        
        # 청크들을 임베딩하고 벡터스토어에 저장
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        documents = []
        embeddings = []
        metadatas = []
        
        for i, chunk in enumerate(chunks):
            status_text.text(f"임베딩 생성 중... ({i+1}/{len(chunks)})")
            progress_bar.progress((i + 1) / len(chunks))
            
            # Gemini로 임베딩 생성
            embedding = get_gemini_embedding(chunk['text'], api_key)
            if embedding:
                documents.append(chunk['text'])
                embeddings.append(embedding)
                metadatas.append({
                    'start': chunk['start'],
                    'end': chunk['end'],
                    'page': 'guide_2022',
                    'chunk_id': i
                })
        
        # 벡터스토어 생성
        vector_store = {
            'documents': documents,
            'embeddings': np.array(embeddings),
            'metadatas': metadatas,
            'created_at': datetime.now().isoformat()
        }
        
        # pickle로 저장
        if documents:
            with open(vector_store_path, 'wb') as f:
                pickle.dump(vector_store, f)
            st.success(f"✅ {len(documents)}개 청크를 벡터스토어에 저장했습니다.")
        
        progress_bar.empty()
        status_text.empty()
        
        return vector_store
        
    except Exception as e:
        st.error(f"벡터스토어 생성 오류: {str(e)}")
        return None

def is_valid_text(text):
    """텍스트 품질 검사"""
    if not text or len(text.strip()) < 10:
        return False

    # 한글 깨짐 검사 (깨진 문자 비율이 30% 이상이면 제외)
    broken_chars = sum(1 for char in text if ord(char) > 55000)  # 한글 깨짐 문자 범위
    if len(text) > 0 and broken_chars / len(text) > 0.3:
        return False

    # 점선 과다 검사 (점선이 50% 이상이면 제외)
    dot_chars = text.count('·') + text.count('…') + text.count('.')
    if len(text) > 0 and dot_chars / len(text) > 0.5:
        return False

    # 반복 문자 과다 검사
    import re
    repeated_patterns = re.findall(r'(.)\1{10,}', text)  # 같은 문자가 10번 이상 반복
    if repeated_patterns:
        return False

    return True

def clean_text_content(text):
    """텍스트 정제"""
    import re

    # 1. 과도한 점선 제거
    text = re.sub(r'[·…]{3,}', ' ', text)
    text = re.sub(r'\.{3,}', ' ', text)

    # 2. 과도한 공백 정리
    text = re.sub(r'\s+', ' ', text)

    # 3. 페이지 번호 패턴 제거
    text = re.sub(r'\b\d+\s*페이지?\b', '', text)
    text = re.sub(r'\b\d+\s*쪽?\b', '', text)

    # 4. 목차 관련 패턴 제거
    text = re.sub(r'^[IVX]+\.?\s*', '', text, flags=re.MULTILINE)  # 로마숫자
    text = re.sub(r'^\d+\.?\s*$', '', text, flags=re.MULTILINE)   # 단독 숫자

    # 5. 반복되는 특수문자 정리
    text = re.sub(r'[~`!@#$%^&*()_+=\[\]{}|\\:";\'<>?/,-]{5,}', ' ', text)

    return text.strip()

def extract_legal_reasoning_from_analysis(analysis_text):
    """Gemini 분석 결과에서 법적 근거와 논리 추출"""
    import re

    extracted_context = {
        'legal_basis': [],      # 법적 근거 (법령, 조항)
        'reasoning': [],        # 추론 과정
        'key_concepts': [],     # 핵심 개념
        'problem_details': []   # 구체적인 문제점
    }

    # 1. 법령 및 조항 추출
    legal_references = re.findall(r'(?:지방자치법|헌법|행정기본법|건축법|도시계획법)\s*(?:제\s*\d+조?(?:의?\d+)?)?', analysis_text)
    extracted_context['legal_basis'].extend(legal_references)

    # 2. 법적 원칙/개념 추출
    legal_concepts = [
        '기관위임사무', '자치사무', '국가사무', '법률유보원칙', '권한배분',
        '상위법령', '법령우위', '조례제정권', '위임입법', '처분권한',
        '헌법위반', '기본권침해', '평등원칙', '비례원칙', '신뢰보호',
        '재산권침해', '영업의자유', '거주이전의자유', '표현의자유',
        '조세법률주의', '죄형법정주의', '적법절차', '정당한보상'
    ]

    for concept in legal_concepts:
        if concept in analysis_text:
            # 해당 개념 주변 문맥 추출 (앞뒤 50자)
            matches = re.finditer(re.escape(concept), analysis_text)
            for match in matches:
                start = max(0, match.start() - 50)
                end = min(len(analysis_text), match.end() + 50)
                context = analysis_text[start:end].strip()
                extracted_context['key_concepts'].append({
                    'concept': concept,
                    'context': context
                })

    # 3. 문제점 상세 내용 추출
    problem_patterns = [
        r'문제(?:점|가|는)[^.]*?(?:\.|$)',
        r'위법[^.]*?(?:\.|$)',
        r'위반[^.]*?(?:\.|$)',
        r'부적절[^.]*?(?:\.|$)',
        r'한계[^.]*?(?:\.|$)'
    ]

    for pattern in problem_patterns:
        matches = re.findall(pattern, analysis_text, re.DOTALL)
        extracted_context['problem_details'].extend(matches)

    # 4. 추론 과정 추출 (따라서, 그러므로, 왜냐하면 등)
    reasoning_patterns = [
        r'(?:따라서|그러므로|왜냐하면|이는|이에 따라)[^.]*?(?:\.|$)',
        r'(?:근거|이유|원인)는[^.]*?(?:\.|$)'
    ]

    for pattern in reasoning_patterns:
        matches = re.findall(pattern, analysis_text, re.DOTALL)
        extracted_context['reasoning'].extend(matches)

    # 5. 디버깅용 출력

    return extracted_context

def search_precedents(query_keywords, max_results=10):
    """국가법령정보센터 API를 통한 판례 검색"""
    try:
        # 🆕 검색 키워드 최적화 ('조례' 단독 검색 우선)
        # 먼저 '조례'만으로 검색을 시도하고, 필요시 개별 키워드 추가 검색
        search_query = "조례"

        # API 요청 파라미터
        params = {
            'OC': OC,
            'target': 'prec',  # 판례 검색
            'type': 'XML',
            'query': search_query,
            'display': min(max_results, 20)  # 최대 20개
        }

        st.info(f"🔍 판례 검색 중: '{search_query}'")

        response = requests.get(precedent_search_url, params=params, timeout=30)
        if response.status_code != 200:
            st.warning(f"판례 검색 API 오류: HTTP {response.status_code}")
            return []

        root = ET.fromstring(response.text)
        precedents = []

        # 🆕 XML 응답 파싱 (올바른 구조 사용)
        # 루트가 PrecSearch이므로, prec 태그를 직접 찾음
        for prec_elem in root.findall('prec'):
            try:
                prec_id = prec_elem.find('판례일련번호')
                case_name = prec_elem.find('사건명')
                court = prec_elem.find('법원명')
                date = prec_elem.find('선고일자')
                case_type = prec_elem.find('사건종류명')

                if all(elem is not None for elem in [prec_id, case_name]):
                    precedent = {
                        'id': prec_id.text,
                        'case_name': case_name.text,
                        'court': court.text if court is not None else '',
                        'date': date.text if date is not None else '',
                        'case_type': case_type.text if case_type is not None else '',
                        'summary': ''  # 요약문은 상세 조회에서 가져올 예정
                    }
                    precedents.append(precedent)
            except Exception as e:
                continue

        st.success(f"📋 {len(precedents)}개의 관련 판례를 발견했습니다.")

        # 🆕 추가 키워드별 검색 (OR 방식 구현)
        if isinstance(query_keywords, list) and len(query_keywords) > 0:
            st.info("🔄 각 키워드별로 관련 판례를 추가 검색합니다...")

            # 각 키워드별로 개별 검색 수행
            additional_precedents = []
            search_keywords = [k for k in query_keywords[:3] if k.strip()]  # 빈 문자열 제거

            for keyword in search_keywords:
                try:
                    # '조례 + 키워드' 조합으로 검색
                    combined_query = f"조례 {keyword}"
                    st.info(f"🔍 키워드별 검색: '{combined_query}'")

                    keyword_params = params.copy()
                    keyword_params['query'] = combined_query
                    keyword_params['display'] = 3  # 각 키워드당 3개씩

                    keyword_response = requests.get(precedent_search_url, params=keyword_params, timeout=15)
                    if keyword_response.status_code == 200:
                        keyword_root = ET.fromstring(keyword_response.text)
                        keyword_precs = keyword_root.findall('prec')

                        st.info(f"   → '{keyword}' 키워드로 {len(keyword_precs)}개 판례 발견")

                        for prec_elem in keyword_precs:
                            try:
                                prec_id = prec_elem.find('판례일련번호')
                                case_name = prec_elem.find('사건명')

                                if prec_id is not None and case_name is not None:
                                    # 중복 제거
                                    all_existing_ids = [p['id'] for p in precedents + additional_precedents]

                                    if prec_id.text not in all_existing_ids:
                                        court = prec_elem.find('법원명')
                                        date = prec_elem.find('선고일자')
                                        case_type = prec_elem.find('사건종류명')

                                        additional_precedent = {
                                            'id': prec_id.text,
                                            'case_name': case_name.text,
                                            'court': court.text if court is not None else '',
                                            'date': date.text if date is not None else '',
                                            'case_type': case_type.text if case_type is not None else '',
                                            'summary': ''
                                        }
                                        additional_precedents.append(additional_precedent)

                            except Exception as e:
                                continue

                except Exception as e:
                    st.warning(f"키워드 '{keyword}' 검색 중 오류: {str(e)}")
                    continue

            if additional_precedents:
                st.success(f"✅ 추가로 {len(additional_precedents)}개의 판례를 더 발견했습니다!")
                precedents.extend(additional_precedents)
            else:
                st.info("추가 검색에서는 새로운 판례를 찾지 못했습니다.")

        return precedents[:max_results]

    except Exception as e:
        st.error(f"판례 검색 오류: {str(e)}")
        return []

def get_precedent_detail(precedent_id):
    """판례 상세 내용 조회"""
    try:
        params = {
            'OC': OC,
            'target': 'prec',
            'ID': precedent_id,
            'type': 'XML'
        }

        response = requests.get(detail_url, params=params, timeout=30)
        if response.status_code != 200:
            return None

        root = ET.fromstring(response.text)

        # 판례 본문 추출
        content = ""

        # 판시사항
        decision_matters = root.find('.//판시사항')
        if decision_matters is not None and decision_matters.text:
            content += f"[판시사항]\n{decision_matters.text}\n\n"

        # 판결요지
        decision_summary = root.find('.//판결요지')
        if decision_summary is not None and decision_summary.text:
            content += f"[판결요지]\n{decision_summary.text}\n\n"

        # 참조조문
        ref_articles = root.find('.//참조조문')
        if ref_articles is not None and ref_articles.text:
            content += f"[참조조문]\n{ref_articles.text}\n\n"

        # 전문 (주요 부분만)
        full_text = root.find('.//전문')
        if full_text is not None and full_text.text:
            # 전문이 너무 길 경우 앞부분만 가져옴
            full_content = full_text.text
            if len(full_content) > 2000:
                full_content = full_content[:2000] + "..."
            content += f"[전문]\n{full_content}\n\n"

        return content.strip() if content else None

    except Exception as e:
        st.warning(f"판례 상세 조회 오류: {str(e)}")
        return None

def extract_legal_principles_from_precedents(precedents_content):
    """판례에서 법리 추출"""
    legal_principles = []

    for i, content in enumerate(precedents_content):
        if not content:
            continue

        # 법리 추출 패턴
        principle_patterns = [
            r'법원은.*?고 판시하였다',
            r'대법원은.*?고 본다',
            r'이 사건에 관하여.*?것이다',
            r'따라서.*?할 것이다',
            r'그러므로.*?라고 할 것이다',
            r'헌법재판소는.*?고 판단한다'
        ]

        extracted_principles = []
        for pattern in principle_patterns:
            matches = re.findall(pattern, content, re.DOTALL)
            for match in matches:
                # 문장 정리
                clean_principle = re.sub(r'\s+', ' ', match.strip())
                if len(clean_principle) > 50 and clean_principle not in extracted_principles:
                    extracted_principles.append(clean_principle)

        if extracted_principles:
            legal_principles.extend(extracted_principles[:2])  # 판례당 최대 2개 법리

    return legal_principles[:6]  # 전체 최대 6개 법리

def search_relevant_guidelines(query, vector_store, api_key=None, top_k=3):
    """쿼리와 관련된 가이드라인 검색 (Gemini 기반 또는 무료 버전)"""
    try:

        if not vector_store or 'embeddings' not in vector_store:
            return []


        # 벡터스토어 타입 확인 (Gemini 기반 vs 무료 버전)
        is_free_version = 'model_name' in vector_store and isinstance(vector_store['model_name'], str)

        if is_free_version:
            # 무료 sentence-transformers 기반 검색
            try:
                from sentence_transformers import SentenceTransformer
                model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
                query_embedding = model.encode([query])[0]
            except ImportError:
                st.warning("sentence-transformers 라이브러리가 필요합니다.")
                return []
        else:
            # Gemini 기반 검색
            if not api_key:
                return []
            query_embedding = get_gemini_embedding(query, api_key)
            if not query_embedding:
                return []

        # 코사인 유사도 계산
        query_embedding = np.array(query_embedding).reshape(1, -1)
        similarities = cosine_similarity(query_embedding, vector_store['embeddings'])[0]

        # 최소 유사도 필터링 (기준 상향)
        min_similarity = 0.5 if is_free_version else 0.3  # 기준을 높여서 관련성 높은 결과만
        valid_indices = np.where(similarities >= min_similarity)[0]

        # 추가적으로 키워드 기반 관련성 검사 (개선된 버전)
        keyword_filtered_indices = []
        law_keywords = ['조례', '법률', '규정', '위반', '위법', '허가', '승인', '사무', '권한', '기관위임', '재의', '제소', '의결', '대법원', '판례']

        for idx in valid_indices:
            try:
                text = vector_store['documents'][idx]
                # 한글이 포함되어 있고 의미 있는 내용인지 확인
                korean_chars = sum(1 for char in text if '\uac00' <= char <= '\ud7af')

                if korean_chars >= 10:  # 최소 10개 이상의 한글이 있어야 함
                    text_lower = text.lower()
                    keyword_count = sum(1 for keyword in law_keywords if keyword in text_lower)

                    # 조건 완화: 1개 이상의 키워드만 있어도 포함
                    if keyword_count >= 1:
                        keyword_filtered_indices.append(idx)
                    else:
                        pass  # 키워드 부족한 경우 (디버깅 코드 제거됨)
                else:
                    pass  # 한글 부족한 경우 (디버깅 코드 제거됨)

            except Exception as e:
                continue

        if keyword_filtered_indices:
            valid_indices = np.array(keyword_filtered_indices)
        else:
            pass  # 키워드 필터링 결과가 없는 경우 (디버깅 코드 제거됨)

        if len(valid_indices) == 0:
            return []

        # 상위 k개 결과 선택
        valid_similarities = similarities[valid_indices]
        top_indices = valid_indices[np.argsort(valid_similarities)[-top_k:][::-1]]
        
        relevant_chunks = []
        for idx in top_indices:
            original_text = vector_store['documents'][idx]

            # 텍스트 품질 검사 및 필터링
            if not is_valid_text(original_text):
                continue

            # 텍스트 정제
            cleaned_text = clean_text_content(original_text)

            if len(cleaned_text.strip()) < 50:  # 너무 짧은 텍스트 제외
                continue

            relevant_chunks.append({
                'text': cleaned_text,
                'original_text': original_text,  # 디버깅용 원본 보관
                'similarity': similarities[idx],
                'distance': 1 - similarities[idx],
                'metadata': vector_store['metadatas'][idx],
                'source': vector_store.get('pdf_path', 'unknown')
            })
        
        return relevant_chunks
        
    except Exception as e:
        st.error(f"가이드라인 검색 오류: {str(e)}")
        return []

def search_multiple_vectorstores(query, api_key=None, top_k_per_store=2):
    """여러 벡터스토어에서 복합 검색을 수행하는 함수"""
    try:
        # 사용 가능한 벡터스토어 경로들
        vectorstore_paths = [
            "enhanced_vectorstore_20250914_101739.pkl",  # 향상된 벡터스토어 (양쪽 PDF 포함, 리랭커 지원)
        ]
        
        vectorstore_names = {
            "enhanced_vectorstore_20250914_101739.pkl": "통합 법령 문서 (재의·제소 + 자치법규입안가이드)"
        }
        
        all_results = []
        loaded_stores = []
        
        for path in vectorstore_paths:
            if os.path.exists(path):
                try:
                    with open(path, 'rb') as f:
                        vector_store = pickle.load(f)
                    
                    # 각 벡터스토어에서 검색 수행
                    results = search_relevant_guidelines(query, vector_store, api_key, top_k_per_store)
                    
                    # 결과에 소스 정보 추가
                    store_name = vectorstore_names.get(os.path.basename(path), path)
                    for result in results:
                        result['source_store'] = store_name
                        result['source_file'] = path
                    
                    all_results.extend(results)
                    loaded_stores.append(store_name)
                    
                except Exception as e:
                    st.warning(f"{path} 로드 실패: {str(e)}")
                    continue
        
        if not all_results:
            return [], []
        
        # 유사도 기준으로 정렬하고 상위 결과 선택
        all_results.sort(key=lambda x: x['similarity'], reverse=True)
        
        # 최대 6개 결과 반환 (각 스토어당 최대 2개씩)
        final_results = []
        store_counts = {}
        
        for result in all_results:
            store_name = result['source_store']
            if store_counts.get(store_name, 0) < top_k_per_store and len(final_results) < 6:
                final_results.append(result)
                store_counts[store_name] = store_counts.get(store_name, 0) + 1
        
        return final_results, loaded_stores
        
    except Exception as e:
        st.error(f"복합 벡터스토어 검색 오류: {str(e)}")
        return [], []

def extract_legality_keywords_from_analysis(analysis_result, api_key):
    """Gemini 1차 분석 결과에서 위법성 의심 키워드 추출"""
    try:
        if not analysis_result or not api_key:
            return []

        # 키워드 추출을 위한 프롬프트
        keyword_extraction_prompt = f"""
다음은 조례 위법성 분석 결과입니다. 이 분석 결과에서 판례 검색에 유용한 핵심 키워드를 추출해주세요.

**분석 결과**:
{analysis_result[:2000]}  # 토큰 제한을 위해 앞부분만

**추출 조건**:
1. 위법성이 의심되는 구체적인 법적 쟁점 키워드 (예: "기관위임사무", "법정위임한계", "포괄위임금지원칙")
2. 관련 법령이나 제도 키워드 (예: "건축허가", "개발행위허가", "환경영향평가")
3. 판례에서 다뤄질 가능성이 높은 키워드 우선

**출력 형식**: 키워드1, 키워드2, 키워드3 (최대 5개, 쉼표로 구분)

키워드:"""

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash-lite')

        response = model.generate_content(keyword_extraction_prompt)

        if response and hasattr(response, 'text') and response.text:
            # 키워드 파싱
            keywords_text = response.text.strip()
            # "키워드:" 이후 텍스트 추출
            if "키워드:" in keywords_text:
                keywords_text = keywords_text.split("키워드:")[-1].strip()

            # 쉼표로 분리하고 정리
            keywords = [kw.strip() for kw in keywords_text.split(',') if kw.strip()]
            # 불필요한 문자 제거
            cleaned_keywords = []
            for kw in keywords[:5]:  # 최대 5개
                clean_kw = re.sub(r'[^\w가-힣\s]', '', kw).strip()
                if len(clean_kw) >= 2 and clean_kw not in cleaned_keywords:
                    cleaned_keywords.append(clean_kw)

            return cleaned_keywords

        return []

    except Exception as e:
        st.warning(f"키워드 추출 오류: {str(e)}")
        return []

def perform_preliminary_analysis(pdf_text, superior_laws_content, search_results, api_key):
    """1차 예비 분석 수행 - 위법성 의심 사유 파악"""
    try:
        if not api_key:
            return None, []

        # 1차 분석용 간단한 프롬프트 (판례 없이)
        preliminary_prompt = f"""
다음 조례를 분석하여 위법성이 의심되는 핵심 쟁점을 파악해주세요.

**조례 내용**:
{pdf_text[:3000]}

**상위법령 정보**:
{str(superior_laws_content)[:2000] if superior_laws_content else '없음'}

**관련 가이드라인**:
{str(search_results)[:1000] if search_results else '없음'}

**분석 요청**:
1. 가장 심각한 위법성 의심 사유 3개를 간략히 제시
2. 각 사유별로 관련 법적 쟁점 키워드 제시
3. 판례 검색이 필요한 핵심 키워드 추출

**출력 형식**:
## 위법성 의심 사유
1. [사유1]: [구체적 내용]
2. [사유2]: [구체적 내용]
3. [사유3]: [구체적 내용]

## 판례 검색 키워드
[키워드1, 키워드2, 키워드3, 키워드4, 키워드5]
"""

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash-lite')

        response = model.generate_content(preliminary_prompt)

        if response and hasattr(response, 'text') and response.text:
            analysis_text = response.text

            # 키워드 추출
            keywords = extract_legality_keywords_from_analysis(analysis_text, api_key)

            return analysis_text, keywords

        return None, []

    except Exception as e:
        st.error(f"1차 예비 분석 오류: {str(e)}")
        return None, []

def detect_agency_delegation(superior_article: Dict, ordinance_article: Dict, source_type: str) -> Dict:
    """기관위임사무 특화 판별 함수"""
    
    superior_content = superior_article.get('content', '').lower()
    ordinance_content = ordinance_article.get('content', '').lower()
    
    # 1단계: 국가사무인지 판별
    national_affairs_indicators = [
        '건축허가', '개발행위허가', '환경영향평가', '도시계획',
        '산업단지', '관광단지', '택지개발', '도로개설',
        '하천점용', '산지전용', '농지전용', '산업입지',
        '국토계획', '지역계획', '광역계획'
    ]
    
    is_national_affair = any(indicator in superior_content for indicator in national_affairs_indicators)
    
    # 2단계: 지방자치단체 '장'에게 위임되었는지 확인
    delegation_to_head_indicators = [
        '시장', '군수', '구청장', '지방자치단체의 장',
        '시장이', '군수가', '구청장이', '장이',
        '위임한다', '위탁한다'
    ]
    
    is_delegated_to_head = any(indicator in superior_content for indicator in delegation_to_head_indicators)
    
    # 3단계: 조례가 해당 사무에 대해 별도 규정을 두고 있는지 확인
    ordinance_regulation_indicators = [
        '허가', '승인', '신고', '인가', '지정', '등록',
        '기준', '절차', '방법', '조건', '제한'
    ]
    
    has_ordinance_regulation = any(indicator in ordinance_content for indicator in ordinance_regulation_indicators)
    
    # 4단계: 위법성 판단
    is_agency_delegation = False
    severity = "낮음"
    evidence = []
    description = ""
    
    if is_national_affair and is_delegated_to_head and has_ordinance_regulation:
        is_agency_delegation = True
        severity = "매우 높음"
        description = "기관위임사무에 대해 조례로 별도 규정을 두어 지방자치법 제22조 위반"
        
        evidence.extend([
            f"국가사무 확인: {[ind for ind in national_affairs_indicators if ind in superior_content][:2]}",
            f"지방자치단체 장 위임 확인: {[ind for ind in delegation_to_head_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정 확인: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    elif is_national_affair and has_ordinance_regulation:
        # 국가사무인데 조례로 규정한 경우 (위임 대상 불확실)
        is_agency_delegation = True
        severity = "높음"
        description = "국가사무로 추정되는 사항에 대해 조례가 별도 규정, 기관위임사무 가능성"
        
        evidence.extend([
            f"국가사무 가능성: {[ind for ind in national_affairs_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    elif is_delegated_to_head and has_ordinance_regulation:
        # 지방자치단체 장 위임 + 조례 규정
        is_agency_delegation = True
        severity = "높음" 
        description = "지방자치단체 장에게 위임된 사무에 대해 조례로 별도 규정"
        
        evidence.extend([
            f"지방자치단체 장 위임: {[ind for ind in delegation_to_head_indicators if ind in superior_content][:2]}",
            f"조례 별도 규정: {[ind for ind in ordinance_regulation_indicators if ind in ordinance_content][:2]}"
        ])
    
    return {
        'is_agency_delegation': is_agency_delegation,
        'description': description,
        'evidence': evidence,
        'severity': severity,
        'national_affair': is_national_affair,
        'delegated_to_head': is_delegated_to_head,
        'has_regulation': has_ordinance_regulation
    }

def analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content):
    """조례와 상위법령 직접 비교 분석 함수 - 계층별 통합 검토"""
    analysis_results = []
    
    if not superior_laws_content:
        return "상위법령 정보가 없어 직접 비교 분석을 수행할 수 없습니다."
    
    
    # 조례에서 사무 관련 조문 추출
    ordinance_provisions = []
    lines = pdf_text.split('\n')
    current_article = ""
    current_content = ""
    
    for line in lines:
        line = line.strip()
        if line.startswith('제') and '조' in line:
            if current_article:
                ordinance_provisions.append({
                    'article': current_article,
                    'content': current_content.strip()
                })
            current_article = line
            current_content = ""
        else:
            current_content += line + " "
    
    # 마지막 조문 추가
    if current_article:
        ordinance_provisions.append({
            'article': current_article,
            'content': current_content.strip()
        })
    
    # 상위법령과 직접 비교 분석
    comparison_results = []
    
    for ordinance_provision in ordinance_provisions:
        if not ordinance_provision['content']:
            continue
            
        provision_analysis = {
            'ordinance_article': ordinance_provision['article'],
            'ordinance_content': ordinance_provision['content'],
            'superior_law_conflicts': [],
            'delegation_issues': [],
            'authority_issues': []
        }
        
        # 각 상위법령 그룹과 비교 (법률, 시행령, 시행규칙 통합)
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            
            # 연결된 본문이 있는 경우 간단한 키워드 매칭만 수행
            if 'combined_content' in law_group:
                superior_content_lower = law_group['combined_content'].lower()
                ordinance_lower = ordinance_provision['content'].lower()
                
                # 키워드 기반 관련성 확인
                common_keywords = []
                for word in ordinance_lower.split():
                    if len(word) > 2 and word in superior_content_lower:
                        common_keywords.append(word)
                
                if len(common_keywords) > 2:  # 최소 3개 이상의 공통 키워드가 있으면 관련성 있음
                    # 간단한 분석만 수행
                    continue
                else:
                    continue
            
            # 기존 방식 - articles가 있는 경우
            for superior_article in law_group.get('combined_articles', []):
                superior_content = superior_article['content'].lower()
                ordinance_lower = ordinance_provision['content'].lower()
                
                # 어느 계층(법률/시행령/시행규칙)에서 나온 조문인지 확인
                article_source = "법률"  # 기본값
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        for article in law_info['articles']:
                            if article['content'] == superior_article['content']:
                                if law_type == 'law':
                                    article_source = "법률"
                                elif law_type == 'decree':
                                    article_source = "시행령"
                                elif law_type == 'rule':
                                    article_source = "시행규칙"
                                break
                
                # 🆕 특화된 기관위임사무 판별 로직
                agency_delegation_result = detect_agency_delegation(
                    superior_article, ordinance_article, article_source
                )
                
                if agency_delegation_result['is_agency_delegation']:
                    provision_analysis['delegation_issues'].append({
                        'superior_law': f"{base_name} ({article_source})",
                        'superior_article': f"{superior_article['number']} {superior_article['title']}",
                        'superior_content': superior_article['content'],
                        'issue_type': '기관위임사무 위반',
                        'description': agency_delegation_result['description'],
                        'evidence': agency_delegation_result['evidence'],
                        'severity': agency_delegation_result['severity'],
                        'hierarchy': article_source
                    })
                
                # 직접적인 충돌 검사 - 계층별 위반 심각도 구분
                conflict_indicators = [
                    ('금지', '허용'), ('의무', '면제'), ('필수', '선택'),
                    ('강제', '임의'), ('반드시', '가능'), ('불가', '허용')
                ]
                
                for prohibit_word, allow_word in conflict_indicators:
                    if prohibit_word in superior_content and allow_word in ordinance_lower:
                        # 계층별 위반 심각도
                        severity = "심각" if article_source == "법률" else ("보통" if article_source == "시행령" else "경미")
                        
                        provision_analysis['superior_law_conflicts'].append({
                            'superior_law': f"{base_name} ({article_source})",
                            'superior_article': f"{superior_article['number']} {superior_article['title']}",
                            'conflict_type': f'{article_source} {prohibit_word} vs 조례 {allow_word}',
                            'superior_content': superior_article['content'],
                            'potential_violation': True,
                            'hierarchy': article_source,
                            'severity': severity
                        })
        
        if provision_analysis['delegation_issues'] or provision_analysis['superior_law_conflicts']:
            comparison_results.append(provision_analysis)
    
    return comparison_results

def create_analysis_prompt(pdf_text, search_results, superior_laws_content=None, relevant_guidelines=None, is_first_ordinance=False, comprehensive_analysis_results=None, theoretical_results=None, precedents_content=None, legal_principles=None):
    """분석 프롬프트 생성 함수"""
    prompt = (
        "🚨 **핵심 미션: 근거 기반 위법성 판정**\n"
        "너는 조례 위법성 전문 검토관이다. 제공된 **구체적 근거 자료들을 활용하여** 위법 여부를 판정하는 것이 목표다.\n\n"

        "**🔍 분석 방법론:**\n"
        "1. **PKL 검색 결과 활용**: 제공된 위법 판례와 현재 조례의 유사성 분석\n"
        "2. **판례 검색 결과 활용**: 국가법령정보센터에서 검색된 관련 판례의 법리 적용\n"
        "3. **상위법령 직접 비교**: 조례 조문과 상위법령 조문의 구체적 대조 분석\n"
        "4. **가이드라인 참조**: 자치법규 작성 가이드라인의 검토 기준 적용\n\n"

        "**📋 작성 원칙:**\n"
        "- ❌ 금지: '~하다면 위법이다', '~할 경우 문제가 된다' 등의 가정적 표현\n"
        "- ✅ 필수: '조례 제○조 \"(조문 인용)\"는 ○○법 제○조 \"(조문 인용)\"와 다음과 같이 충돌한다'\n"
        "- ✅ 필수: PKL/판례 검색 결과에서 발견된 유사 사례와 현재 조례의 구체적 비교\n"
        "- ✅ 필수: 검색된 관련 판례의 법리를 현재 조례에 직접 적용한 분석\n"
        "- ✅ 필수: 위법이 없으면 '검토 결과 위법 사항 없음'으로 명확히 결론\n\n"

        "**📄 현재 검토 대상 조례 전문:**\n"
        "---\n"
        f"{pdf_text}\n"
        "---\n"
    )
    
    # 상위법령 내용 추가 (계층별 그룹화)
    if superior_laws_content:
        prompt += "\n그리고 아래는 조례안에서 언급된 상위법령들의 실제 조문 내용이야. (법률, 시행령, 시행규칙을 계층별로 그룹화하여 통합 분석)\n"
        prompt += "---\n"
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            prompt += f"◆ {base_name}\n"
            
            # 연결된 본문이 있으면 사용
            if 'combined_content' in law_group:
                prompt += f"  본문 내용:\n{law_group['combined_content']}\n"
            else:
                # 기존 방식 - 각 계층별 법령 표시
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                        prompt += f"  [{type_name}] {law_info['law_name']}\n"
                
                # 통합된 조문 표시 (상위 15개만)
                prompt += f"  통합 조문 ({len(law_group['combined_articles'])}개):\n"
                for article in law_group['combined_articles'][:15]:  
                    prompt += f"    {article['number']} {article['title']}\n"
                    prompt += f"    {article['content']}\n\n"
        prompt += "---\n"
        
        # 상위법령 직접 비교 분석 결과 추가
        try:
            comparison_results = analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content)
            if comparison_results and isinstance(comparison_results, list) and len(comparison_results) > 0:
                prompt += "\n**중요: 조례와 상위법령 직접 비교 분석 결과**\n"
                prompt += "아래는 조례 조문과 상위법령을 하나씩 직접 비교한 결과이다. 이 분석을 바탕으로 기관위임사무 여부와 법령위반 가능성을 정확히 판단해줘.\n"
                prompt += "---\n"
                
                for result in comparison_results:
                    prompt += f"◆ {result['ordinance_article']}\n"
                    prompt += f"조례 내용: {result['ordinance_content'][:200]}...\n"
                    
                    if result['delegation_issues']:
                        prompt += "⚠️ 기관위임사무 가능성 발견:\n"
                        for issue in result['delegation_issues']:
                            prompt += f"  - {issue['superior_law']} {issue['superior_article']}\n"
                            prompt += f"    문제: {issue['description']}\n"
                    
                    if result['superior_law_conflicts']:
                        prompt += "🚨 상위법령 충돌 가능성 발견:\n"
                        for conflict in result['superior_law_conflicts']:
                            prompt += f"  - {conflict['superior_law']} {conflict['superior_article']}\n"
                            prompt += f"    충돌: {conflict['conflict_type']}\n"
                    
                    prompt += "\n"
                prompt += "---\n"
        except Exception as e:
            prompt += f"\n상위법령 직접 비교 분석 중 오류 발생: {str(e)}\n"
    
    # 자치법규 가이드라인 및 사례 추가
    if relevant_guidelines:
        prompt += "\n**📋 통합 법령 문서 검색 결과 (재의제소 + 자치법규입안가이드)**\n"
        prompt += "자치법규 작성 가이드라인과 과거 문제 사례들이다.\n"
        prompt += "**분석 방법**: 아래 가이드라인과 사례를 현재 조례와 직접 비교하여 위법 여부를 판정하라.\n"
        prompt += "'가이드라인에서 ○○는 금지한다고 했는데, 현재 조례 제○조가 이에 해당한다' 식으로 구체적 지적하라.\n"
        prompt += "---\n"
        
        # 소스별로 그룹화하여 표시
        source_groups = {}
        for guideline in relevant_guidelines:
            source_store = guideline.get('source_store', '알 수 없는 자료')
            if source_store not in source_groups:
                source_groups[source_store] = []
            source_groups[source_store].append(guideline)
        
        for source_store, guidelines in source_groups.items():
            prompt += f"◆ 참고자료: {source_store}\n"
            for i, guideline in enumerate(guidelines):
                similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                prompt += f"  [{i+1}] (유사도: {similarity_score:.3f})\n"
                prompt += f"  {guideline['text']}\n\n"
        prompt += "---\n"
    
    # 종합 위법성 판례 분석 결과 추가
    if comprehensive_analysis_results and isinstance(comprehensive_analysis_results, list) and len(comprehensive_analysis_results) > 0:
        total_risks = sum(len(result['violation_risks']) for result in comprehensive_analysis_results)
        prompt += f"\n**📊 PKL 검색 결과 기반 분석 ({total_risks}개 위험 발견)**\n"
        prompt += "아래는 PKL 파일에서 검색된 실제 조례 위법 판례들을 현재 조례에 적용한 분석 결과이다.\n"
        prompt += "**중요**: 각 판례의 위법 사유와 현재 조례 조문을 직접 비교하여 위법 여부를 구체적으로 판정하라.\n"
        prompt += "단순히 '유사하므로 위법 가능성이 있다'가 아니라, '어떤 부분이 어떻게 위법인지' 명확히 지적하라.\n"
        prompt += "---\n"
        
        for result in comprehensive_analysis_results:
            prompt += f"◆ {result['ordinance_article']}\n"
            prompt += f"조례 내용: {result['ordinance_content'][:150]}...\n"
            
            for i, risk in enumerate(result['violation_risks'][:2]):  # 상위 2개만 포함
                prompt += f"  위험 {i+1}: {risk['violation_type']} (위험도: {risk['risk_score']:.2f}/1.0)\n"
                prompt += f"  관련 판례: {risk['case_summary'][:150]}...\n"
                if risk['legal_principle'] != "해당없음":
                    prompt += f"  법적 원칙: {risk['legal_principle']}\n"
                prompt += f"  개선 권고: {risk['recommendation']}\n"
                prompt += f"  판례 출처: {risk['case_source']}\n\n"
            
            if len(result['violation_risks']) > 2:
                prompt += f"  ...외 {len(result['violation_risks']) - 2}개 추가 위험\n\n"
        prompt += "---\n"

    # 🆕 검색된 관련 판례/이론 추가
    if theoretical_results and isinstance(theoretical_results, list) and len(theoretical_results) > 0:
        prompt += f"\n**📚 PKL 추가 검색 결과 분석 ({len(theoretical_results)}개 관련 자료)**\n"
        prompt += "1차 분석에서 발견된 문제점들과 관련된 추가 판례와 법리이다.\n"
        prompt += "**분석 방법**: 각 자료의 내용과 현재 조례를 구체적으로 대조하여 위법 여부를 판정하라.\n"
        prompt += "가설이나 추정이 아닌, 실제 조문 비교를 통한 명확한 결론을 제시하라.\n"
        prompt += "---\n"

        for i, theory in enumerate(theoretical_results[:5]):  # 상위 5개만 포함
            context_rel = theory.get('context_relevance', 0)
            matched_concepts = theory.get('matched_concepts', [])
            similarity = theory.get('similarity', 0)

            prompt += f"◆ 관련 판례/이론 {i+1} (관련도: {context_rel:.2f}, 유사도: {similarity:.2f})\n"
            if matched_concepts:
                prompt += f"관련 개념: {', '.join(matched_concepts)}\n"

            # 내용 미리보기 (300자로 제한)
            content = theory.get('content', theory.get('text', '내용 없음'))
            content_preview = content[:300] + "..." if len(content) > 300 else content
            prompt += f"내용: {content_preview}\n\n"

        prompt += "**⚠️ 중요**: 위 판례들은 조례의 문제점과 직접 관련이 있으므로, 이를 근거로 현재 조례의 위법성을 구체적으로 지적하고 개선방안을 제시하라.\n"
        prompt += "---\n"

    # 판례 검색 결과 추가
    if precedents_content and len(precedents_content) > 0:
        prompt += f"\n**⚖️ 국가법령정보센터 판례 검색 결과 ({len(precedents_content)}개)**\n"
        prompt += "조례 관련 쟁점에 대한 판례들이다. **중요**: 각 판례의 법리를 현재 조례의 구체적 조문에 적용하라.\n"
        prompt += "**분석 방법**: '판례에서 ○○는 위법하다고 했는데, 현재 조례 제○조도 동일한 내용이므로 위법이다' 식으로 구체적 비교하라.\n"
        prompt += "---\n"

        for i, precedent in enumerate(precedents_content[:3]):  # 최대 3개 판례
            prompt += f"◆ 판례 {i+1}\n"
            if isinstance(precedent, dict):
                if 'case_name' in precedent:
                    prompt += f"사건명: {precedent['case_name']}\n"
                if 'court' in precedent:
                    prompt += f"법원: {precedent['court']}\n"
                if 'date' in precedent:
                    prompt += f"선고일: {precedent['date']}\n"
                content = precedent.get('content', '')
            else:
                content = str(precedent)

            # 판례 내용 요약 (500자 제한)
            if len(content) > 500:
                content = content[:500] + "..."
            prompt += f"판례 내용:\n{content}\n\n"

        prompt += "---\n"

    # 추출된 법리 추가
    if legal_principles and len(legal_principles) > 0:
        prompt += f"\n**📖 판례로부터 추출된 법리 ({len(legal_principles)}개)**\n"
        prompt += "위 판례들로부터 추출된 핵심 법리들이다. 이 법리들을 현재 조례에 적용하여 위법성을 구체적으로 판단하라.\n"
        prompt += "---\n"

        for i, principle in enumerate(legal_principles[:5]):  # 최대 5개 법리
            prompt += f"{i+1}. {principle}\n\n"

        prompt += "**📍 중요**: 위 법리들을 근거로 현재 조례의 구체적인 조문이 어떤 법적 문제가 있는지 명확히 지적하고, 개선방안을 제시하라.\n"
        prompt += "---\n"

    if is_first_ordinance:
        prompt += (
            "※ 참고: 이 조례는 17개 시도 중 최초로 제정되는 조례로, 타시도 조례가 존재하지 않습니다.\n"
            "타시도 조례가 없는 상황에서, 아래 기준에 따라 조례의 적정성, 상위법령과의 관계, 실무적 검토 포인트 등을 중심으로 분석해줘.\n"
        )
    else:
        prompt += "그리고 아래는 타시도 조례명과 각 조문 내용이야.\n"
        for result in search_results:
            prompt += f"조례명: {result['name']}\n"
            for idx, article in enumerate(result['content']):
                prompt += f"제{idx+1}조: {article}\n"
    
    prompt += (
        "---\n"
        "**🎯 최종 분석 지시사항**\n"
        "위에 제공된 모든 검색 결과(PKL, 판례, 가이드라인, 상위법령)를 종합하여 다음과 같이 분석하라.\n\n"

        "**📊 분석 방법론:**\n"
        "1. **근거 자료 우선**: 검색된 판례와 가이드라인을 구체적 근거로 활용\n"
        "2. **조문 대조**: 현재 조례 조문과 상위법령/판례를 직접 비교\n"
        "3. **명확한 결론**: '위법 사항 있음' 또는 '위법 사항 없음'으로 명확히 결론\n"
        "4. **구체적 지적**: 가설이 아닌 실제 비교를 통한 위법 지적\n\n"

        "이제 아래 기준에 따라 분석해줘. 반드시 한글로 답변해줘.\n"
        "1. [근거 기반 위법성 분석]\n"
        "- 위에 제공된 PKL 검색 결과, 판례 검색 결과, 가이드라인을 활용한 구체적 분석\n"
        "- '○○ 판례에서 금지한 ○○○가 현재 조례 제○조에 동일하게 나타남' 식으로 구체적 지적\n\n"
        "2. [비교분석 요약표(조문별)]\n"
        "- 표의 컬럼: 조문(내 조례), 주요 내용, 타 시도 유사 조항, 동일 여부, 차이 및 내 조례 특징, 추천 조문\n"
        "- 반드시 내 조례(PDF로 업로드한 조례)의 조문만을 기준으로, 각 조문별로 타 시도 조례와 비교해 표로 정리(내 조례에 없는 조문은 비교하지 말 것)\n"
        "- '추천 조문' 칸에는 타 시도 조례와 비교해 무난하게 생각되는 조문 예시를 한글로 작성\n\n"
        "3. [내 조례의 차별점 요약] (별도 소제목)\n"
        "- 타 시도 조례와 비교해 독특하거나 구조적으로 다른 점, 내 조례만의 관리/운영 방식 등 요약\n\n"
        "4. [검토 시 유의사항] (별도 소제목)\n"
        "각 항목마다 일반인도 이해할 수 있도록 쉬운 말로 부연설명도 함께 작성해줘.\n"
        "다음 원칙들을 기준으로 검토해줘:\n"
        "a) 소관사무의 원칙 - **🚨 매우 중요: 기관위임사무는 조례 제정 금지**\n"
        "**기관위임사무 정의**: 국가사무를 지방자치단체의 '장'(시장, 군수, 구청장)에게 위임한 사무\n"
        "**핵심 원칙**: 기관위임사무에 대해서는 조례 제정이 원칙적으로 금지됨 (지방자치법 제22조)\n"
        "**판별 기준**: \n"
        "  1) 사무가 국가사무인지 확인 (예: 건축허가, 도시계획, 환경영향평가 등)\n"
        "  2) 해당 사무가 지방자치단체 '장'에게 위임되었는지 확인\n"
        "  3) 위임된 사무에 대해 조례가 별도 규정을 두고 있는지 검토\n"
        "**위법 사례**: 건축허가, 개발행위허가, 환경영향평가 등 국가위임사무에 대해 조례로 추가 규정을 둔 경우\n"
        "- 지방자치단체의 자치사무와 법령에 의해 위임된 단체위임사무에 대해서만 제정 가능한지\n"
        "- 사무의 성격이 전국적으로 통일적 처리를 요구하는지 여부 검토\n\n"
        "b) 법률 유보의 원칙\n"
        "- 주민의 권리를 제한하거나 의무를 부과하는 내용이 있는지\n"
        "- 상위 법령에서 위임받지 않은 권한을 행사하는지\n"
        "- 상위 법령의 위임 범위를 초과하는지\n\n"
        "c) 법령우위의 원칙 위반 여부 \n"
        "- **🚨 매우 중요: 실제 위법 내용을 찾아내는 것이 목표**\n"
        "- **일반론이 아닌 구체적 충돌 지점을 반드시 찾아라**\n"
        "- 위에 제시된 상위법령 본문을 한 조문씩 꼼꼼히 읽고 조례와 직접 대조하라\n\n"
        "**검토 방법**:\n"
        "1) 조례 제1조부터 마지막 조문까지 하나씩 검토\n"
        "2) 각 조례 조문의 내용과 관련된 상위법령 조문을 찾아서 직접 비교\n"
        "3) 다음과 같은 구체적 충돌이 있는지 확인:\n"
        "   - 조례가 금지하는 것을 상위법령이 허용하는 경우\n"
        "   - 조례가 허용하는 것을 상위법령이 금지하는 경우\n"
        "   - 조례가 상위법령보다 강한 의무나 제재를 부과하는 경우\n"
        "   - 조례가 상위법령의 위임 범위를 명백히 벗어나는 경우\n"
        "   - 조례가 상위법령에서 국가나 중앙행정기관 소관으로 정한 사무에 관여하는 경우\n\n"
        "**위법 발견 시 반드시 다음 형식으로 구체적으로 명시:**\n"
        "  🚨 **위법 사항 발견**\n"
        "  * **조례 조문**: 제○조 ○항 - \"조례의 정확한 문구\"\n"
        "  * **상위법령**: ○○법 제○조 ○항 - \"상위법령의 정확한 문구\"\n"
        "  * **충돌 내용**: 구체적으로 어떤 부분이 어떻게 위배되는지 상세 설명\n"
        "  * **위법 유형**: (법령우위 위반/법률유보 위반/기관위임사무 위반)\n"
        "  * **개선 방안**: 상위법령에 맞는 구체적 수정안\n\n"
        "**위법 사항이 없는 경우에만** '위법 사항을 발견하지 못했음'이라고 결론짓고,\n"
        "**의심스러운 부분이 있으면 반드시 지적**하라.\n\n"
        "4. 실무적 검토 포인트\n"
        "- 조례의 집행 과정에서 발생할 수 있는 문제점\n"
        "- 개선이 필요한 부분과 그 방향성\n\n"
    )

    # 상위법령별 개별 위반 여부 검토 (Gemini 전용 프롬프트 추가)
    if superior_laws_content:
        prompt += "\n5. [상위법령별 개별 위반 여부 검토]\n"
        prompt += "위에서 제시한 상위법령들 각각에 대해 개별적으로 다음 기준에 따라 상세 분석해줘:\n\n"

        section_num = 1
        for law_group in superior_laws_content:
            base_name = law_group['base_name']

            prompt += f"5-{section_num}) [{base_name} 위반 여부 검토]\n"
            prompt += f"상위법령명: {base_name}\n"

            # 해당 법령의 본문 일부 재참조
            if 'combined_content' in law_group:
                law_content_preview = law_group['combined_content'][:2000]
                prompt += f"상위 법령 본문 일부:\n{law_content_preview}\n\n"
            elif 'combined_articles' in law_group and law_group['combined_articles']:
                prompt += "상위 법령 주요 조문:\n"
                for article in law_group['combined_articles'][:5]:  # 처음 5개 조문만
                    prompt += f"  {article['number']} {article['title']}\n"
                    prompt += f"  {article['content'][:300]}...\n\n"

            prompt += f"**🔍 {base_name} 세부 검토 지시사항:**\n"
            prompt += "위 상위법령 본문을 조례와 한 조문씩 직접 대조하여 다음을 수행하라:\n\n"
            prompt += "  ① **조문별 직접 대조 분석**\n"
            prompt += f"  - 조례의 각 조문이 {base_name}의 어떤 조문과 관련되는지 식별\n"
            prompt += f"  - {base_name}에서 금지/허용/의무화하는 사항과 조례 내용 직접 비교\n"
            prompt += "  - 상충되는 부분이 있으면 구체적으로 지적\n\n"
            prompt += "  ② **권한 범위 초과 여부**\n"
            prompt += f"  - {base_name}에서 국가/중앙행정기관 전담으로 정한 사무가 있는지 확인\n"
            prompt += "  - 조례가 해당 사무에 개입하고 있는지 점검\n"
            prompt += "  - 위임 범위를 벗어난 규정이 있는지 확인\n\n"
            prompt += "  ③ **구체적 위법 사항 발견 시**\n"
            prompt += "  🚨 **위법 발견 보고 형식:**\n"
            prompt += "  * **문제 조문**: 조례 제○조 - \"정확한 조문 내용\"\n"
            prompt += f"  * **관련 상위법령**: {base_name} 제○조 - \"정확한 조문 내용\"\n"
            prompt += "  * **위법 사유**: 구체적인 충돌/위반 내용\n"
            prompt += "  * **위법 심각도**: 경미/보통/심각\n"
            prompt += "  * **수정 방안**: 구체적인 개선 방향\n\n"
            prompt += "  ④ **의심 사항도 반드시 보고**\n"
            prompt += "  - 명확하지 않지만 위법 가능성이 있는 부분\n"
            prompt += "  - 해석에 따라 문제가 될 수 있는 조문\n\n"

            section_num += 1

    return prompt

def parse_table_from_text(text_content):
    """텍스트에서 표 형태의 내용을 파싱하여 Word 표 데이터로 변환"""
    tables_data = []
    lines = text_content.split('\n')
    current_table = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 표의 시작을 감지 (|가 포함된 라인)
        if '|' in line and len([cell for cell in line.split('|') if cell.strip()]) >= 3:
            # 표 헤더인지 구분 (첫 번째 |로 시작하는 라인)
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]

            if current_table is None:
                # 새 표 시작
                current_table = {'headers': cells, 'rows': []}
                tables_data.append(current_table)
            else:
                # 구분선이 아닌 데이터 행인지 확인
                if not all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                    current_table['rows'].append(cells)
        else:
            # 표가 끝남
            if current_table is not None:
                current_table = None

    return tables_data

def add_table_to_doc(doc, table_data):
    """Word 문서에 표 추가"""
    if not table_data['headers']:
        return

    # 열 수 계산
    max_cols = len(table_data['headers'])
    for row in table_data['rows']:
        max_cols = max(max_cols, len(row))

    # 행 수 계산 (헤더 + 데이터 행)
    row_count = 1 + len(table_data['rows'])

    if row_count == 1:  # 헤더만 있는 경우 스킵
        return

    # 표 생성
    table = doc.add_table(rows=row_count, cols=max_cols)
    table.style = 'Table Grid'
    table.autofit = True

    # 헤더 추가
    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data['headers']):
        if i < len(header_cells):
            header_cells[i].text = header
            # 헤더 스타일링
            paragraph = header_cells[i].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.bold = True

    # 데이터 행 추가
    for row_idx, row_data in enumerate(table_data['rows']):
        if row_idx + 1 < len(table.rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = cell_data

def create_comparison_document(pdf_text, search_results, analysis_results, superior_laws_content=None, relevant_guidelines=None):
    """비교 분석 문서 생성 함수"""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)

    # 제목 추가
    title = doc.add_heading('조례 비교 분석 결과', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'분석 일시: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')

    # 상위법령 정보 추가 (계층별 그룹화)
    if superior_laws_content:
        doc.add_heading('검토된 상위법령', level=2)
        for law_group in superior_laws_content:
            base_name = law_group['base_name']
            
            # 그룹 제목 추가
            doc.add_paragraph(f"◆ {base_name}")
            
            # 연결된 본문이 있는 경우
            if 'combined_content' in law_group:
                content_length = len(law_group['combined_content'])
                doc.add_paragraph(f"  • 본문 {content_length:,}자")
            else:
                # 기존 방식 - 각 계층별 법령 정보 표시
                for law_type, law_info in law_group['laws'].items():
                    if law_info and 'articles' in law_info:
                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                        doc.add_paragraph(f"  • {law_info['law_name']} ({type_name}) - {len(law_info['articles'])}개 조문")
                
                combined_articles = law_group.get('combined_articles', [])
                doc.add_paragraph(f"  총 {len(combined_articles)}개 조문 통합 검토")
            
            doc.add_paragraph("")
        doc.add_paragraph("")
    
    # 활용된 자치법규 자료 정보 추가
    if relevant_guidelines:
        doc.add_heading('활용된 자치법규 참고자료', level=2)
        
        # 소스별로 그룹화
        source_groups = {}
        for guideline in relevant_guidelines:
            source_store = guideline.get('source_store', '알 수 없는 자료')
            if source_store not in source_groups:
                source_groups[source_store] = []
            source_groups[source_store].append(guideline)
        
        for source_store, guidelines in source_groups.items():
            doc.add_paragraph(f"◆ {source_store} ({len(guidelines)}개 내용)")
            for i, guideline in enumerate(guidelines):
                similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                doc.add_paragraph(f"   • 내용 {i+1} (유사도: {similarity_score:.3f})")
        doc.add_paragraph("")

    # 각 API 분석 결과 추가
    for result in analysis_results:
        if 'error' in result:
            doc.add_paragraph(f"{result['model']} API 오류: {result['error']}")
            continue

        content = result['content']

        # 🆕 표 파싱 및 처리
        tables_data = parse_table_from_text(content)

        # 텍스트를 섹션별로 처리
        lines = content.split('\n')
        current_section = []

        for line in lines:
            line = line.strip()

            # 표 라인인지 확인 (|가 포함된 라인)
            if '|' in line and len([cell for cell in line.split('|') if cell.strip()]) >= 3:
                # 표 시작 전까지의 텍스트 처리
                if current_section:
                    for text_line in current_section:
                        text_line_clean = text_line.strip()
                        if text_line_clean:
                            # 제목 라인 처리 (1., 2., 3. 등으로 시작하거나 [로 시작하는 경우)
                            if (text_line_clean.startswith(('1.', '2.', '3.', '4.', '5.')) or
                                text_line_clean.startswith('[') and text_line_clean.endswith(']')):
                                # 마크다운 기호 제거하고 제목으로 추가
                                title_text = re.sub(r'[#*`>\-\[\]]+', '', text_line_clean)
                                doc.add_heading(title_text, level=3)
                            else:
                                # 일반 텍스트
                                clean_text = re.sub(r'[#*`>]+', '', text_line_clean)
                                if clean_text:
                                    doc.add_paragraph(clean_text)
                    current_section = []

                # 표 처리는 skip (이미 tables_data에서 처리됨)
                continue
            else:
                # 구분선이 아닌 경우만 텍스트로 추가
                if not (line.replace('-', '').replace(':', '').replace('|', '').strip() == ''):
                    current_section.append(line)

        # 마지막 섹션 처리
        if current_section:
            for text_line in current_section:
                text_line_clean = text_line.strip()
                if text_line_clean:
                    if (text_line_clean.startswith(('1.', '2.', '3.', '4.', '5.')) or
                        text_line_clean.startswith('[') and text_line_clean.endswith(']')):
                        title_text = re.sub(r'[#*`>\-\[\]]+', '', text_line_clean)
                        doc.add_heading(title_text, level=3)
                    else:
                        clean_text = re.sub(r'[#*`>]+', '', text_line_clean)
                        if clean_text:
                            doc.add_paragraph(clean_text)

        # 🆕 파싱된 표들을 Word 문서에 추가
        for table_data in tables_data:
            add_table_to_doc(doc, table_data)
            doc.add_paragraph("")  # 표 간격

    return doc

def main():
    # 헤더
    st.markdown("""
    <div class="main-header">
        <h1>🏛️ 광역지자체 조례 검색, 비교, 분석</h1>
        <p>17개 광역지자체의 조례를 검색하고, AI를 활용하여 비교 분석할 수 있는 도구입니다.</p>
    </div>
    """, unsafe_allow_html=True)

    # 사이드바
    with st.sidebar:
        st.header("📋 작업 순서")
        st.markdown("""
        <div class="step-card">
            <strong>1단계:</strong> 조례 검색 및 Word 저장<br>
            검색어를 입력하여 17개 시도의 조례를 검색하고 3단 비교 형태로 MS Word 문서를 생성합니다.
        </div>
        <div class="step-card">
            <strong>2단계:</strong> 조례안 PDF 업로드<br>
            제정 또는 개정할 조례안 PDF 파일을 업로드합니다.
        </div>
        <div class="step-card">
            <strong>3단계:</strong> AI 비교 분석<br>
            업로드한 조례안과 타 시도 조례를 AI로 비교 분석하여 MS Word 문서를 생성합니다.
        </div>
        """, unsafe_allow_html=True)

        st.header("🔑 API 설정")
        gemini_api_key = st.text_input("Gemini API 키", type="password", help="Google AI Studio에서 발급받은 API 키를 입력하세요")
        openai_api_key = st.text_input("OpenAI API 키", type="password", help="OpenAI 플랫폼에서 발급받은 API 키를 입력하세요")
        
        st.header("🔑 API 키 설정 가이드")
        st.markdown("""
        <div class="step-card">
            <strong>📋 API 키 발급 및 설정 방법</strong><br>
            조례 분석을 위한 AI 서비스 API 키를 발급받고 설정하는 방법을 안내합니다.
        </div>
        """, unsafe_allow_html=True)
        
        # 🆕 상세한 API 키 설정 가이드
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("### 🤖 Gemini API 키 발급")
            with st.expander("📋 단계별 가이드", expanded=False):
                st.markdown("""
                **1. Google AI Studio 접속**
                - 브라우저에서 [aistudio.google.com](https://aistudio.google.com) 접속
                - Google 계정으로 로그인

                **2. API 키 생성**
                - 좌측 메뉴에서 'API Keys' 클릭
                - 'Create API Key' 버튼 클릭
                - 프로젝트 선택 (없으면 새로 생성)

                **3. API 키 복사**
                - 생성된 API 키를 복사
                - 안전한 곳에 보관 (재확인 불가)

                **4. 사용량 확인**
                - 무료 할당량: 월 1,000번 요청
                - 유료 전환 시 더 많은 사용량 제공

                ⚠️ **주의**: API 키는 개인정보이므로 타인과 공유하지 마세요!
                """)

        with col2:
            st.markdown("### 🧠 OpenAI API 키 발급")
            with st.expander("📋 단계별 가이드", expanded=False):
                st.markdown("""
                **1. OpenAI 플랫폼 접속**
                - 브라우저에서 [platform.openai.com](https://platform.openai.com) 접속
                - OpenAI 계정 생성/로그인

                **2. API 키 생성**
                - 우상단 프로필 → 'API keys' 클릭
                - 'Create new secret key' 버튼 클릭
                - 키 이름 입력 후 생성

                **3. 결제 정보 등록**
                - 'Billing' 메뉴에서 결제수단 등록
                - 사용량 한도 설정 (권장: $10-20)

                **4. 요금 정보**
                - GPT-4: 입력 토큰당 $0.03/1K, 출력 토큰당 $0.06/1K
                - 일반적으로 분석 1회당 $0.5-2 정도 소요

                💡 **팁**: 처음에는 낮은 한도로 시작하여 사용량을 확인해보세요.
                """)

        # 벡터스토어 자동 로드 (백그라운드에서 조용히 처리)
        vector_store_path = "enhanced_vectorstore_20250914_101739.pkl"
        if st.session_state.vector_store is None and os.path.exists(vector_store_path):
            try:
                with open(vector_store_path, 'rb') as f:
                    st.session_state.vector_store = pickle.load(f)
            except Exception:
                pass  # 조용히 실패

    # 메인 컨텐츠
    tab1, tab2, tab3 = st.tabs(["1️⃣ 조례 검색", "2️⃣ PDF 업로드", "3️⃣ AI 분석"])

    with tab1:
        st.header("조례 검색")
        
        # 검색 폼 (Enter 키 지원)
        with st.form(key="search_form"):
            col1, col2 = st.columns([3, 1])
            with col1:
                search_query = st.text_input(
                    "검색어를 입력하세요 (키워드)", 
                    placeholder="예: 청년지원 (Enter 키로도 검색 가능)", 
                    value=st.session_state.search_query,
                    help="검색어를 입력한 후 Enter 키를 누르거나 검색 버튼을 클릭하세요."
                )
            with col2:
                search_button = st.form_submit_button("🔍 검색", type="primary")

        # 검색 실행 (Enter 키 또는 버튼 클릭 시)
        if search_button and search_query.strip():
            st.session_state.search_query = search_query.strip()
            st.session_state.word_doc_ready = False  # 문서 준비 상태 초기화
            st.session_state.selected_ordinances = []  # 선택된 조례 초기화
            
            with st.spinner("검색 중... 잠시만 기다려주세요."):
                try:
                    results, total_count = search_ordinances(search_query.strip())
                    st.session_state.search_results = results
                    # 초기에는 모든 조례를 선택된 상태로 설정
                    st.session_state.selected_ordinances = list(range(len(results)))
                    st.success(f"검색 완료! 총 {len(results)}건의 조례가 검색되었습니다.")
                except Exception as e:
                    st.error(f"검색 중 오류 발생: {str(e)}")
                    st.session_state.search_results = []

        # 검색 결과가 있을 때 조례 선택 및 Word 문서 생성 기능
        if st.session_state.search_results:
            results = st.session_state.search_results
            
            # 검색 결과 요약 표시
            if not st.session_state.word_doc_ready:
                st.success(f"검색 완료! 총 {len(results)}건의 조례가 검색되었습니다.")
            
            # 조례 선택 섹션
            st.subheader("📋 Word 문서에 포함할 조례 선택")
            
            # 전체 선택/해제 버튼
            col1, col2, col3 = st.columns([1, 1, 2])
            
            with col1:
                if st.button("✅ 전체 선택", key="select_all_btn"):
                    st.session_state.selected_ordinances = list(range(len(results)))
                    st.rerun()
            
            with col2:
                if st.button("❌ 전체 해제", key="deselect_all_btn"):
                    st.session_state.selected_ordinances = []
                    st.rerun()
            
            with col3:
                selected_count = len(st.session_state.selected_ordinances)
                st.markdown(f"**선택된 조례: {selected_count}개 / 총 {len(results)}개**")
            
            # 조례 선택 체크박스
            st.markdown("---")
            
            # 조례별 체크박스 표시
            for idx, result in enumerate(results):
                # 🆕 단순화: 체크박스 상태를 직접 관리
                is_selected = idx in st.session_state.selected_ordinances
                checkbox_key = f"ordinance_checkbox_{idx}"

                # 체크박스와 조례명을 한 줄에 표시
                current_checked = st.checkbox(
                    f"**{result['metro']}** - {result['name']}",
                    value=is_selected,
                    key=checkbox_key
                )

                # 🆕 상태 변경 감지 및 즉시 반영
                if current_checked != is_selected:
                    if current_checked:
                        # 체크됨 - 목록에 추가
                        if idx not in st.session_state.selected_ordinances:
                            st.session_state.selected_ordinances.append(idx)
                    else:
                        # 체크 해제됨 - 목록에서 제거
                        if idx in st.session_state.selected_ordinances:
                            st.session_state.selected_ordinances.remove(idx)
            
            st.markdown("---")
            
            # Word 문서 생성 버튼
            col1, col2 = st.columns([1, 1])
            
            with col1:
                # 선택된 조례가 있을 때만 생성 버튼 활성화
                disabled = len(st.session_state.selected_ordinances) == 0
                
                if st.button("📄 선택된 조례로 Word 문서 생성", type="secondary", key="create_word_btn", disabled=disabled):
                    if st.session_state.selected_ordinances:
                        try:
                            with st.spinner("Word 문서 생성 중..."):
                                # 선택된 조례만 필터링
                                selected_results = [results[i] for i in st.session_state.selected_ordinances]
                                
                                # Word 문서 생성
                                doc = create_word_document(st.session_state.search_query, selected_results)
                                
                                # Word 문서를 바이트로 변환
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                doc_bytes = doc_io.getvalue()
                                
                                # 세션 상태에 저장
                                st.session_state.word_doc_data = doc_bytes
                                st.session_state.word_doc_ready = True
                                
                            st.success(f"✅ 선택된 {len(selected_results)}개 조례로 Word 문서가 생성되었습니다!")
                            st.rerun()  # 페이지 새로고침으로 다운로드 버튼 표시
                            
                        except Exception as e:
                            st.error(f"❌ Word 문서 생성 중 오류 발생: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())
                    else:
                        st.warning("조례를 하나 이상 선택해주세요.")
                
                if disabled:
                    st.caption("⚠️ 조례를 하나 이상 선택해주세요.")
            
            with col2:
                # Word 문서가 준비되면 다운로드 버튼 표시
                if st.session_state.word_doc_ready and st.session_state.word_doc_data:
                    filename = f"조례_검색결과_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.download_button(
                        label="💾 Word 문서 다운로드",
                        data=st.session_state.word_doc_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_word_btn"
                    )
            # 상세 검색 결과 표시 (조례 내용 확인용)
            st.subheader("📖 조례 내용 상세보기")
            
            for idx, result in enumerate(results):
                # 🆕 단순화: 선택 상태만 텍스트로 표시
                is_selected = idx in st.session_state.selected_ordinances
                status = " ✅ 선택됨" if is_selected else " ⭕ 선택안됨"

                with st.expander(f"{result['metro']} - {result['name']}{status}", expanded=False):
                    st.markdown(f"<div class='metro-name'>{result['metro']}</div>", unsafe_allow_html=True)
                    st.markdown(f"<div class='law-title'>{result['name']}</div>", unsafe_allow_html=True)
                    
                    if result['content']:
                        for article_idx, article in enumerate(result['content']):
                            st.markdown(f"**제{article_idx+1}조**")
                            st.markdown(article)
                            st.markdown("---")
                    else:
                        st.markdown("*(조문 없음)*")
        
        elif search_button and not search_query.strip():
            st.error("검색어를 입력해주세요.")
        elif not st.session_state.search_results:
            st.info("검색어를 입력하고 Enter 키를 누르거나 검색 버튼을 클릭하세요.")

    with tab2:
        st.header("조례안 PDF 업로드")
        
        uploaded_file = st.file_uploader("제정 또는 개정할 조례안 PDF 파일을 업로드하세요", type=['pdf'])
        
        if uploaded_file is not None:
            st.session_state.uploaded_pdf = uploaded_file
            st.success(f"파일이 업로드되었습니다: {uploaded_file.name}")
            
            # PDF 내용 미리보기
            if st.checkbox("PDF 내용 미리보기"):
                with st.spinner("PDF 내용을 읽는 중..."):
                    pdf_text = extract_pdf_text(uploaded_file)
                    if pdf_text:
                        st.text_area("PDF 내용", pdf_text[:2000] + "..." if len(pdf_text) > 2000 else pdf_text, height=300)
                    else:
                        st.error("PDF 내용을 읽을 수 없습니다.")

    with tab3:
        st.header("AI 비교 분석")
        
        # 조건 확인 - PDF가 업로드되고 API 키가 있으면 분석 가능
        pdf_uploaded = st.session_state.uploaded_pdf is not None
        has_api_key = bool(gemini_api_key or openai_api_key)
        has_search_results = bool(st.session_state.search_results)
        
        if not pdf_uploaded:
            st.warning("📄 먼저 PDF 파일을 업로드해주세요.")
        elif not has_api_key:
            st.warning("🔑 API 키를 하나 이상 입력해주세요.")
        else:
            # 검색 결과 여부에 따라 안내 메시지 표시
            if not has_search_results:
                st.info("💡 **최초 제정 조례 분석**")
                st.markdown("""
                검색된 타 시도 조례가 없습니다. 이는 다음과 같은 경우일 수 있습니다:
                - 🆕 **최초 제정 조례**: 17개 시도 중 최초로 제정되는 조례
                - 🔍 **검색어 불일치**: 다른 키워드로 재검색 후 분석 권장
                
                검색 결과가 없어도 조례안의 **법적 검토**와 **상위법령 위반 여부** 분석이 가능합니다.
                """)
            else:
                st.success(f"📊 {len(st.session_state.search_results)}개의 타 시도 조례와 비교 분석합니다.")
        
        # 분석 가능한 조건일 때 분석 인터페이스 표시
        if pdf_uploaded and has_api_key:
            # 검색어 입력 (선택사항)
            search_query_analysis = st.text_input(
                "검색어 (분석용)", 
                value=st.session_state.search_query if st.session_state.search_query else "", 
                key="analysis_query",
                help="검색어를 입력하면 더 정확한 분석이 가능합니다. (선택사항)"
            )
            
            # 분석 타입 표시 (선택된 조례 수 반영)
            if not has_search_results:
                analysis_type = "최초 제정 조례 분석"
            elif hasattr(st.session_state, 'selected_ordinances') and st.session_state.selected_ordinances:
                selected_count = len(st.session_state.selected_ordinances)
                analysis_type = f"선택된 {selected_count}개 타 시도 조례와 비교 분석"
            else:
                analysis_type = f"전체 {len(st.session_state.search_results)}개 타 시도 조례와 비교 분석"
            st.markdown(f"**분석 유형**: {analysis_type}")
            
            # PKL 파일 참고 옵션 (문제 발견 시 자동 활용)
            use_pkl_auto = st.checkbox(
                "🔍 문제 발견 시 PKL 파일 자동 참고", 
                value=True, 
                help="Gemini가 법적 문제점을 발견한 경우 자동으로 무료 벡터스토어를 참고하여 근거를 보강합니다."
            )
            
            # 🆕 저장된 분석 결과가 있으면 먼저 표시
            if hasattr(st.session_state, 'analysis_results') and st.session_state.analysis_results:
                st.info("💾 **이전 분석 결과가 저장되어 있습니다**")

                # 메타데이터 표시
                if hasattr(st.session_state, 'analysis_metadata'):
                    metadata = st.session_state.analysis_metadata
                    st.caption(f"📅 분석 시간: {metadata.get('analysis_timestamp', '알 수 없음')}")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📋 이전 분석 결과 보기", use_container_width=True):
                        st.session_state.show_previous_analysis = True
                        st.rerun()
                with col2:
                    if st.button("🔄 새로 분석하기", use_container_width=True):
                        # 기존 결과 초기화
                        if hasattr(st.session_state, 'analysis_results'):
                            del st.session_state.analysis_results
                        if hasattr(st.session_state, 'analysis_metadata'):
                            del st.session_state.analysis_metadata
                        if hasattr(st.session_state, 'show_previous_analysis'):
                            del st.session_state.show_previous_analysis
                        st.rerun()

            # 이전 분석 결과 표시
            if hasattr(st.session_state, 'show_previous_analysis') and st.session_state.show_previous_analysis and hasattr(st.session_state, 'analysis_results'):
                analysis_results = st.session_state.analysis_results
                metadata = st.session_state.analysis_metadata

                st.markdown("---")
                st.subheader("📋 저장된 AI 분석 결과")

                # 분석 완료 메시지 (저장된 메타데이터 기반)
                has_problems = metadata.get('has_problems', False)
                relevant_guidelines = metadata.get('relevant_guidelines')
                loaded_stores = metadata.get('loaded_stores')
                is_first_ordinance = metadata.get('is_first_ordinance', False)

                if has_problems and relevant_guidelines and loaded_stores:
                    st.success(f"🎯 **복합 자료 보강 분석 완료**: 문제점 탐지 → {len(loaded_stores)}개 자료 참고 → 보강 분석")
                elif has_problems and relevant_guidelines:
                    st.success("🎯 **지능형 분석 완료**: 문제점 탐지 → PKL 참고 → 보강 분석")
                elif has_problems:
                    st.info("⚠️ **문제점 탐지 분석 완료**: PKL 참고 없이 기본 분석만 수행")
                else:
                    st.success("✅ **기본 분석 완료**: 특별한 문제점이 발견되지 않음")

                # 분석 결과 요약
                analysis_count = len([r for r in analysis_results if 'error' not in r])
                if analysis_count > 0:
                    # 🆕 저장된 메타데이터에서 선택된 조례 수 반영
                    if is_first_ordinance:
                        analysis_type_text = "최초 제정 조례"
                    else:
                        saved_search_results = metadata.get('search_results_for_analysis', [])
                        selected_count = len(saved_search_results)
                        analysis_type_text = f"선택된 {selected_count}개 타 시도 조례 비교"
                    st.markdown(f"**📋 분석 유형**: {analysis_type_text}")
                    st.markdown(f"**🤖 수행된 분석**: {analysis_count}개")
                    if relevant_guidelines:
                        guideline_count = len(relevant_guidelines) if isinstance(relevant_guidelines, list) else 0
                        st.markdown(f"**📚 참고 가이드라인**: {guideline_count}개")

                # 분석 결과 표시
                for result in analysis_results:
                    if 'error' not in result:
                        final_report = result
                        # 모델에 따른 구분 표시
                        if "보강" in final_report['model']:
                            st.success("🎯 **복합 자료 참고 보강 분석 결과**")
                            st.caption(f"📚 **활용 모델**: {final_report['model']}")
                        elif "PKL 보강" in final_report['model']:
                            st.success("🎯 **PKL 가이드라인 참고 보강 분석 결과**")
                        elif "OpenAI" in final_report['model']:
                            st.info("📊 **OpenAI 추가 분석 결과**")
                        else:
                            st.info("🤖 **Gemini 기본 분석 결과**")
                        # 보고서 내용
                        st.markdown(final_report['content'])

                # 오류 메시지 표시
                for result in analysis_results:
                    if 'error' in result:
                        st.error(f"❌ {result['model']} 오류: {result['error']}")

                # Word 문서 다운로드 (메타데이터에서 복원)
                with st.spinner("저장된 분석 결과 Word 문서 생성 중..."):
                    superior_laws_content = metadata.get('superior_laws_content')
                    search_results_for_analysis = metadata.get('search_results_for_analysis')
                    pdf_text = metadata.get('pdf_text')
                    doc = create_comparison_document(pdf_text, search_results_for_analysis, analysis_results, superior_laws_content, relevant_guidelines)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_bytes = doc_io.getvalue()
                    # 파일명 설정
                    if has_problems and relevant_guidelines and loaded_stores:
                        stores_count = len(loaded_stores)
                        filename_prefix = f"복합자료보강분석({stores_count}개자료)" if is_first_ordinance else f"조례비교_복합자료분석({stores_count}개자료)"
                    elif has_problems and relevant_guidelines:
                        filename_prefix = "지능형PKL보강분석" if is_first_ordinance else "조례비교_PKL보강분석"
                    elif has_problems:
                        filename_prefix = "문제점탐지분석" if is_first_ordinance else "조례비교_문제점분석"
                    else:
                        filename_prefix = "최초조례_기본분석" if is_first_ordinance else "조례_기본비교분석"
                    st.download_button(
                        label="📄 분석 결과 Word 문서 다운로드",
                        data=doc_bytes,
                        file_name=f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_previous_analysis"
                    )

                st.markdown("---")
                st.markdown("💡 **새로 분석하려면 위의 '🔄 새로 분석하기' 버튼을 클릭하세요.**")

            else:
                # 저장된 결과가 없거나 새 분석을 선택한 경우만 분석 시작 버튼 표시
                # 🆕 선택된 조례가 없는 경우 경고 표시
                if has_search_results and hasattr(st.session_state, 'selected_ordinances') and not st.session_state.selected_ordinances:
                    st.warning("⚠️ 비교할 조례가 선택되지 않았습니다. 조례 검색 탭에서 조례를 선택하거나, 선택 없이 최초 제정 조례 분석을 진행하세요.")

                if st.button("🤖 AI 분석 시작", type="primary"):
                    with st.spinner("AI가 조례를 분석하고 있습니다... 잠시만 기다려주세요."):
                        # PDF 텍스트 추출
                        pdf_text = extract_pdf_text(st.session_state.uploaded_pdf)
                    
                    if not pdf_text:
                        st.error("PDF 텍스트를 읽을 수 없습니다.")
                    else:
                        # 1단계: 상위법령 추출
                        st.info("📋 1단계: 조례안에서 상위법령을 추출하고 있습니다...")
                        superior_laws = extract_superior_laws(pdf_text)
                        
                        if superior_laws:
                            st.success(f"✅ {len(superior_laws)}개의 상위법령을 발견했습니다:")
                            for law in superior_laws:
                                st.markdown(f"   • {law}")
                            
                            # 2단계: 상위법령 내용 조회
                            st.info("📚 2단계: 국가법령정보센터에서 상위법령 내용을 조회하고 있습니다...")
                            superior_laws_content = get_all_superior_laws_content(superior_laws)
                            
                            if superior_laws_content:
                                st.success(f"✅ {len(superior_laws_content)}개의 상위법령 그룹을 성공적으로 조회했습니다:")
                                total_articles = 0
                                for i, law_group in enumerate(superior_laws_content):
                                    base_name = law_group['base_name']
                                    
                                    # 연결된 본문이 있는 경우
                                    if 'combined_content' in law_group:
                                        content_length = len(law_group['combined_content'])
                                        st.markdown(f"   • **{base_name}**: 본문 {content_length:,}자")
                                        total_articles += 1  # 하나의 연결된 법령으로 카운트
                                    else:
                                        # 기존 방식
                                        available_laws = []
                                        group_articles = 0
                                        
                                        for law_type, law_info in law_group['laws'].items():
                                            if law_info:
                                                type_name = {"law": "법률", "decree": "시행령", "rule": "시행규칙"}[law_type]
                                                article_count = len(law_info.get('articles', []))
                                                available_laws.append(f"{type_name}({article_count})")
                                                group_articles += article_count
                                        
                                        st.markdown(f"   • **{base_name}**: {', '.join(available_laws)} = 총 {group_articles}개 조문")
                                        total_articles += group_articles
                                
                                st.markdown(f"   **전체 조문 수**: {total_articles}개")
                                
                                # 🆕 상위법령 본문 내용 디버깅 표시

                                # 연결된 본문이 있는 경우
                                if 'combined_content' in law_group and law_group['combined_content']:
                                    content = law_group['combined_content']
                                    st.markdown(f"**본문 길이**: {len(content):,}자")
                                    st.text_area(
                                        f"{law_group['base_name']} 본문",
                                        content,
                                        height=200,
                                        key=f"content_{i}"
                                    )
                                else:
                                    # 개별 법령별 표시
                                    for law_type, law_info in law_group['laws'].items():
                                        if law_info and 'articles' in law_info:
                                            type_name = {"law": "법률", "decree": "시행령", "rule": "시행규칙"}[law_type]
                                            st.markdown(f"#### {type_name}")

                                            # 조문별 내용 표시 (처음 5개만)
                                            for j, article in enumerate(law_info['articles'][:5]):
                                                st.markdown(f"**제{article.get('number', '?')}조** {article.get('title', '')}")
                                                content = article.get('content', '')[:500]
                                                st.markdown(f"```\n{content}{'...' if len(article.get('content', '')) > 500 else ''}\n```")

                                            if len(law_info['articles']) > 5:
                                                st.markdown(f"... (총 {len(law_info['articles'])}개 조문 중 5개만 표시)")

                                st.markdown("---")
                                
                                # 2-1단계: 상위법령 직접 비교 분석
                                st.info("⚖️ 2-1단계: 조례와 상위법령 직접 비교 분석을 수행합니다...")
                                try:
                                    comparison_results = analyze_ordinance_vs_superior_laws(pdf_text, superior_laws_content)
                                    
                                    if comparison_results and isinstance(comparison_results, list) and len(comparison_results) > 0:
                                        st.warning(f"⚠️ {len(comparison_results)}개 조문에서 잠재적 문제점이 발견되었습니다!")
                                        
                                        with st.expander("🔍 상위법령 직접 비교 분석 결과", expanded=True):
                                            for i, result in enumerate(comparison_results):
                                                st.markdown(f"**🔍 {result['ordinance_article']}**")
                                                st.markdown(f"조례 내용: {result['ordinance_content'][:300]}...")
                                                
                                                if result['delegation_issues']:
                                                    st.error("⚠️ **기관위임사무 가능성 발견**")
                                                    for issue in result['delegation_issues']:
                                                        st.markdown(f"- **관련 상위법령**: {issue['superior_law']} {issue['superior_article']}")
                                                        st.markdown(f"- **문제점**: {issue['description']}")
                                                        st.markdown(f"- **상위법령 내용**: {issue['superior_content'][:200]}...")
                                                
                                                if result['superior_law_conflicts']:
                                                    st.error("🚨 **상위법령 충돌 가능성 발견**")
                                                    for conflict in result['superior_law_conflicts']:
                                                        st.markdown(f"- **관련 상위법령**: {conflict['superior_law']} {conflict['superior_article']}")
                                                        st.markdown(f"- **충돌 유형**: {conflict['conflict_type']}")
                                                        st.markdown(f"- **상위법령 내용**: {conflict['superior_content'][:200]}...")
                                                
                                                st.markdown("---")
                                    else:
                                        st.success("✅ 상위법령 직접 비교에서 명백한 충돌이나 기관위임사무 문제를 발견하지 못했습니다.")
                                        
                                except Exception as e:
                                    st.error(f"상위법령 직접 비교 분석 중 오류: {str(e)}")
                                
                                # 상위법령 내용 미리보기 (계층별 그룹화)
                                    for law_group in superior_laws_content:
                                        base_name = law_group['base_name']
                                        
                                        # 연결된 본문이 있는 경우
                                        if 'combined_content' in law_group:
                                            content_preview = law_group['combined_content'][:500] + "..." if len(law_group['combined_content']) > 500 else law_group['combined_content']
                                            with st.expander(f"📋 {base_name} ({len(law_group['combined_content']):,}자)", expanded=False):
                                                st.text_area("본문 내용", content_preview, height=300, disabled=True)
                                        else:
                                            # 기존 방식
                                            with st.expander(f"📋 {base_name} 계층 ({len(law_group.get('combined_articles', []))}개 조문)", expanded=False):
                                                
                                                # 계층별 법령 정보 표시
                                                st.markdown("**📚 포함된 법령:**")
                                                for law_type, law_info in law_group['laws'].items():
                                                    if law_info and 'articles' in law_info:
                                                        type_name = "법률" if law_type == 'law' else ("시행령" if law_type == 'decree' else "시행규칙")
                                                        st.markdown(f"- [{type_name}] {law_info['law_name']} ({len(law_info['articles'])}개 조문)")
                                                
                                                st.markdown("\n**📖 통합 조문 (처음 5개):**")
                                                combined_articles = law_group.get('combined_articles', [])
                                                for article in combined_articles[:5]:  
                                                    st.markdown(f"**{article['number']} {article['title']}**")
                                                    st.markdown(article['content'][:200] + "..." if len(article['content']) > 200 else article['content'])
                                                    st.markdown("---")
                                                if len(combined_articles) > 5:
                                                    st.markdown(f"*(총 {len(combined_articles)}개 조문 중 5개만 표시)*")
                            else:
                                st.warning("⚠️ 상위법령 내용 조회에 실패했습니다. 일반적인 분석을 진행합니다.")
                        else:
                            st.info("ℹ️ 조례안에서 명시적인 상위법령을 찾을 수 없습니다.")
                            superior_laws_content = None
                        
                        # 3단계: Gemini 1차 분석 (문제점 탐지)
                        st.info("🤖 3단계: Gemini 1차 분석을 수행합니다...")
                        analysis_results = []
                        is_first_ordinance = not has_search_results

                        # 🆕 선택된 조례만 분석에 사용
                        if has_search_results and hasattr(st.session_state, 'selected_ordinances'):
                            selected_results = [st.session_state.search_results[i] for i in st.session_state.selected_ordinances if i < len(st.session_state.search_results)]
                            search_results_for_analysis = selected_results
                            st.info(f"📋 선택된 {len(search_results_for_analysis)}개 조례로 분석을 진행합니다.")
                        else:
                            search_results_for_analysis = st.session_state.search_results if has_search_results else []
                        
                        # Gemini 1차 분석 (문제점 탐지용)
                        first_analysis = None
                        has_problems = False
                        
                        if gemini_api_key:
                            try:
                                # comprehensive_analysis_results 초기화
                                comprehensive_analysis_results = None
                                
                                genai.configure(api_key=gemini_api_key)
                                model = genai.GenerativeModel('gemini-2.0-flash-lite')
                                
                                # 🆕 1차 예비 분석으로 위법성 의심 키워드 추출
                                st.info("🔍 1차 예비 분석: 위법성 의심 사유 파악 중...")
                                preliminary_analysis, legality_keywords = perform_preliminary_analysis(
                                    pdf_text, superior_laws_content, search_results_for_analysis, gemini_api_key
                                )

                                if preliminary_analysis:
                                    st.success("✅ 1차 예비 분석 완료")
                                    with st.expander("🔍 1차 예비 분석 결과 보기"):
                                        st.markdown(preliminary_analysis[:1500] + "..." if len(preliminary_analysis) > 1500 else preliminary_analysis)

                                # 판례 검색은 PKL 분석 이후로 이동
                                precedents = []
                                precedents_content = []
                                legal_principles = []

                                # 🆕 2차 최종 분석용 프롬프트 (1차 분석 + 판례 법리 종합)
                                st.info("📊 2차 최종 분석: 판례 법리 적용한 종합 위법성 판단 중...")
                                theoretical_results = st.session_state.get('theoretical_results', None)
                                guideline_results = st.session_state.get('guideline_results', None)

                                # 기존 함수에 preliminary_analysis를 추가하여 사용
                                final_prompt = create_analysis_prompt(
                                    pdf_text, search_results_for_analysis, superior_laws_content,
                                    guideline_results, is_first_ordinance, comprehensive_analysis_results,
                                    theoretical_results, precedents_content, legal_principles
                                )

                                # 1차 예비 분석 결과를 최종 프롬프트에 추가
                                if preliminary_analysis:
                                    final_prompt = (
                                        "**🔍 1차 예비 분석 결과**\n"
                                        "다음은 위법성 의심 사유를 파악한 1차 예비 분석 결과이다.\n"
                                        "이 결과를 바탕으로 더 정확하고 구체적인 위법성 분석을 수행하라.\n"
                                        "---\n" +
                                        preliminary_analysis +
                                        "\n---\n\n" +
                                        final_prompt
                                    )
                                
                                # 🆕 Gemini 전송 프롬프트 디버깅 표시

                                # 상위법령 내용 부분만 추출
                                if "상위법령들의 실제 조문 내용" in final_prompt:
                                    law_start = final_prompt.find("상위법령들의 실제 조문 내용")
                                    law_end = final_prompt.find("3. [검토 시 유의사항]")
                                    if law_end == -1:
                                        law_end = law_start + 5000  # 기본값

                                    law_content = final_prompt[law_start:law_end]
                                    st.markdown(f"**상위법령 내용 길이**: {len(law_content):,}자")

                                    st.text_area(
                                        "상위법령 관련 프롬프트 내용",
                                                law_content[:3000] + "..." if len(law_content) > 3000 else law_content,
                                                height=300,
                                                key="prompt_law_content"
                                            )
                                # 전체 프롬프트 표시 (처음 2000자만)
                                st.text_area(
                                    "최종 프롬프트 (처음 2000자)",
                                    final_prompt[:2000] + "..." if len(final_prompt) > 2000 else final_prompt,
                                    height=400,
                                    key="final_prompt"
                                )

                                response = model.generate_content(final_prompt)
                                
                                if response and hasattr(response, 'text') and response.text:
                                    first_analysis = response.text
                                    
                                    # 문제점 키워드 탐지
                                    problem_keywords = [
                                        "위반", "문제", "충돌", "부적절", "개선", "수정", "보완",
                                        "법령 위반", "상위법령", "위법", "불일치", "모순", "우려"
                                    ]
                                    
                                    has_problems = any(keyword in first_analysis for keyword in problem_keywords)
                                    
                                    if has_problems:
                                        st.warning(f"⚠️ Gemini가 잠재적 문제점을 발견했습니다!")
                                        
                                        # 🆕 2차 분석: 발견된 문제점에 대한 관련 위법 판례 검색
                                        st.info("🔍 2-0단계: 발견된 문제점에 대한 관련 위법 판례를 PKL에서 검색합니다...")
                                        
                                        try:
                                            from comprehensive_violation_analysis import search_theoretical_background

                                            # 🔍 Gemini 분석 결과에서 구체적인 근거 추출
                                            extracted_context = extract_legal_reasoning_from_analysis(first_analysis)

                                            # 문제점별 이론 검색 (추출된 문맥 활용)
                                            detected_problems = [kw for kw in problem_keywords if kw in first_analysis]
                                            theoretical_results = search_theoretical_background(
                                                detected_problems,
                                                ['3. 지방자치단체의 재의·제소 조례 모음집(Ⅸ) (1)_new_vectorstore.pkl'],
                                                max_results=8,
                                                context_analysis=extracted_context  # 추출된 문맥 전달
                                            )
                                            
                                            if theoretical_results:
                                                st.success(f"✅ {len(theoretical_results)}개의 관련 이론/판례를 찾았습니다!")
                                                
                                                with st.expander("📚 문제점 관련 위법 판례", expanded=False):
                                                    for i, theory in enumerate(theoretical_results[:5]):  # 상위 5개만 표시
                                                        context_rel = theory.get('context_relevance', 0)
                                                        matched_concepts = theory.get('matched_concepts', [])

                                                        st.markdown(f"**[{i+1}] {theory['topic']}**")
                                                        st.markdown(f"📊 **관련도**: {theory['relevance_score']:.3f} | **문맥관련성**: {context_rel}")

                                                        if matched_concepts:
                                                            st.markdown(f"🔍 **매칭된 개념**: {', '.join(matched_concepts[:3])}")

                                                        content_preview = theory['content'][:300] + "..." if len(theory['content']) > 300 else theory['content']
                                                        st.markdown(f"📄 **내용**: {content_preview}")
                                                        st.markdown("---")
                                                
                                                # 위법 판례를 포함한 재분석 프롬프트에 추가할 수 있도록 저장
                                                st.session_state['theoretical_results'] = theoretical_results

                                                # 🆕 통합 법령 문서(재의제소 + 자치법규입안가이드)에서 추가 검색
                                                st.info("📖 통합 법령 문서에서 관련 가이드라인 검색 중...")
                                                try:
                                                    # 추출된 키워드로 통합 법령 문서 검색
                                                    search_query = ' '.join(detected_problems[:3])  # 상위 3개 문제점
                                                    guideline_results, loaded_stores = search_multiple_vectorstores(
                                                        search_query,
                                                        gemini_api_key,
                                                        top_k_per_store=3
                                                    )

                                                    if guideline_results:
                                                        st.success(f"✅ {len(guideline_results)}개의 관련 가이드라인을 찾았습니다!")

                                                        with st.expander("📋 통합 법령 문서 검색 결과", expanded=False):
                                                            for i, result in enumerate(guideline_results):
                                                                st.markdown(f"**[{i+1}] {result.get('source_store', '알 수 없는 출처')}**")
                                                                st.markdown(f"📊 **유사도**: {result.get('similarity', 0):.3f}")

                                                                content_preview = result['text'][:400] + "..." if len(result['text']) > 400 else result['text']
                                                                st.markdown(f"📄 **내용**: {content_preview}")
                                                                st.markdown("---")

                                                        # 세션에 저장
                                                        st.session_state['guideline_results'] = guideline_results
                                                    else:
                                                        st.info("통합 법령 문서에서 관련 내용을 찾지 못했습니다.")
                                                        st.session_state['guideline_results'] = []

                                                except Exception as e:
                                                    st.warning(f"통합 법령 문서 검색 중 오류: {str(e)}")
                                                    st.session_state['guideline_results'] = []

                                            else:
                                                st.warning("관련 위법 판례를 찾지 못했습니다.")
                                                st.session_state['guideline_results'] = []

                                        except Exception as e:
                                            st.error(f"위법 판례 검색 중 오류: {str(e)}")
                                            st.session_state['theoretical_results'] = None

                                        # 🆕 PKL 분석 완료 후 추출된 키워드로 판례 검색
                                        st.info("⚖️ 3단계: 국가법령정보센터에서 관련 판례 검색 중...")

                                        if legality_keywords:
                                            st.info(f"🔎 검색 키워드: {', '.join(legality_keywords)}")

                                            # 위법성 키워드로 판례 검색
                                            search_query = ' '.join(legality_keywords[:3])  # 상위 3개 키워드
                                            precedents = search_precedents(search_query, max_results=5)

                                            if precedents:
                                                st.success(f"📋 {len(precedents)}개 판례 검색 완료")

                                                # 판례 상세 내용 가져오기
                                                progress_bar = st.progress(0)
                                                for i, precedent in enumerate(precedents[:3]):  # 최대 3개만 상세 조회
                                                    detail_content = get_precedent_detail(precedent['id'])
                                                    if detail_content:
                                                        precedent['content'] = detail_content
                                                        precedents_content.append(precedent)
                                                    progress_bar.progress((i+1) / min(len(precedents), 3))

                                                # 판례에서 법리 추출
                                                if precedents_content:
                                                    contents_only = [p.get('content', '') for p in precedents_content]
                                                    legal_principles = extract_legal_principles_from_precedents(contents_only)

                                                    if legal_principles:
                                                        st.success(f"⚖️ {len(legal_principles)}개 법리 추출 완료")
                                                        with st.expander("📖 추출된 법리 보기"):
                                                            for i, principle in enumerate(legal_principles):
                                                                st.markdown(f"{i+1}. {principle}")
                                                    else:
                                                        st.info("법리 추출 결과가 없습니다.")
                                            else:
                                                st.info("관련 판례를 찾을 수 없습니다.")
                                        else:
                                            st.warning("위법성 키워드를 추출할 수 없어 기본 키워드로 검색합니다.")
                                            # 폴백: 조례 제목에서 키워드 추출
                                            fallback_keywords = []
                                            if pdf_text:
                                                title_match = re.search(r'[가-힣\s]{5,30}(?:조례|규칙)', pdf_text[:200])
                                                if title_match:
                                                    title = title_match.group()
                                                    keywords = re.findall(r'[가-힣]{2,6}', title)
                                                    fallback_keywords = keywords[:3]

                                            if fallback_keywords:
                                                search_query = ' '.join(fallback_keywords)
                                                precedents = search_precedents(search_query, max_results=3)

                                                if precedents:
                                                    st.success(f"📋 {len(precedents)}개 판례 검색 완료 (폴백)")
                                                    # 간단한 상세 조회
                                                    for precedent in precedents[:2]:
                                                        detail_content = get_precedent_detail(precedent['id'])
                                                        if detail_content:
                                                            precedent['content'] = detail_content
                                                            precedents_content.append(precedent)
                                                else:
                                                    st.info("폴백 키워드로도 판례를 찾을 수 없습니다.")
                                    else:
                                        st.success("✅ Gemini 1차 분석에서 특별한 문제점이 발견되지 않았습니다.")
                                        
                                    analysis_results.append({
                                        'model': 'Gemini (1차 분석)',
                                        'content': first_analysis
                                    })
                                else:
                                    st.error("Gemini 1차 분석 응답이 비어있습니다.")
                            except Exception as e:
                                st.error(f"Gemini 1차 분석 오류: {str(e)}")
                                analysis_results.append({
                                    'model': 'Gemini (1차 분석)',
                                    'error': str(e)
                                })
                        
                        # 4단계: 문제 발견 시 복합 PKL 참고 분석 수행
                        relevant_guidelines = None
                        loaded_stores = []
                        enhanced_analysis = None
                        
                        if has_problems and use_pkl_auto and first_analysis:
                            st.info("🔍 4단계: 문제점이 발견되어 복합 PKL 파일을 참고한 보강 분석을 수행합니다...")
                            
                            # 🆕 4-1단계: 종합적 조례 위법성 분석 (우선 수행)
                            comprehensive_analysis_results = None
                            if has_problems:  # 문제가 발견된 경우 항상 수행
                                st.info("⚖️ 4-1단계: 모든 유형의 조례 위법 판례를 종합 검색하여 적용합니다...")
                                
                                try:
                                    from comprehensive_violation_analysis import search_comprehensive_violation_cases, apply_violation_cases_to_ordinance
                                    
                                    vectorstore_paths = [
                                        '3. 지방자치단체의 재의·제소 조례 모음집(Ⅸ) (1)_new_vectorstore.pkl'
                                    ]
                                    
                                    # 모든 유형의 위법 사례 종합 검색
                                    # first_analysis에서 조례 정보 추출
                                    ordinance_articles = []
                                    if first_analysis and 'ordinance_data' in first_analysis:
                                        ordinance_articles = first_analysis['ordinance_data']
                                    
                                    violation_cases = search_comprehensive_violation_cases(ordinance_articles, vectorstore_paths, max_results=12)
                                    
                                    if violation_cases:
                                        # 유형별 통계
                                        type_counts = {}
                                        for case in violation_cases:
                                            v_type = case['violation_type']
                                            type_counts[v_type] = type_counts.get(v_type, 0) + 1
                                        
                                        st.success(f"✅ {len(violation_cases)}개의 조례 위법 판례를 발견했습니다:")
                                        
                                        # 유형별 요약
                                        type_summary = []
                                        for v_type, count in type_counts.items():
                                            type_summary.append(f"{v_type} ({count}개)")
                                        st.markdown("**발견된 위법 유형**: " + ", ".join(type_summary))
                                        
                                        # 발견된 판례 미리보기
                                        with st.expander("📚 발견된 조례 위법 판례", expanded=False):
                                            for i, case in enumerate(violation_cases):
                                                st.markdown(f"**[{i+1}] {case['violation_type']}** (유사도: {case['similarity']:.3f})")
                                                st.markdown(f"출처: {case['source_store'].replace('.pkl', '').replace('_', ' ').title()}")
                                                if case['legal_principle'] != "해당없음":
                                                    st.markdown(f"법적 원칙: {case['legal_principle']}")
                                                st.markdown(f"요약: {case['case_summary'][:150]}...")
                                                st.markdown("---")
                                        
                                        # 판례를 현재 조례에 적용하여 종합 위법성 분석
                                        comprehensive_analysis_results = apply_violation_cases_to_ordinance(
                                            violation_cases, pdf_text, superior_laws_content
                                        )
                                        
                                        if comprehensive_analysis_results and isinstance(comprehensive_analysis_results, list):
                                            total_risks = sum(len(result['violation_risks']) for result in comprehensive_analysis_results)
                                            st.warning(f"⚠️ {len(comprehensive_analysis_results)}개 조문에서 총 {total_risks}개의 위법 위험이 발견되었습니다!")
                                            
                                            with st.expander("🚨 종합 위법성 분석 결과", expanded=True):
                                                for result in comprehensive_analysis_results:
                                                    st.error(f"**{result['ordinance_article']}**")
                                                    st.markdown(f"조문 내용: {result['ordinance_content'][:100]}...")
                                                    
                                                    for i, risk in enumerate(result['violation_risks'][:3]):  # 상위 3개만 표시
                                                        st.markdown(f"**위험 {i+1}: {risk['violation_type']}**")
                                                        st.markdown(f"- 위험도: {risk['risk_score']:.2f}/1.0")
                                                        if risk['legal_principle'] != "해당없음":
                                                            st.markdown(f"- 법적 원칙: {risk['legal_principle']}")
                                                        st.markdown(f"- 관련 사례: {risk['case_summary'][:100]}...")
                                                        st.markdown(f"- 개선 권고: {risk['recommendation']}")
                                                        st.markdown("")
                                                    
                                                    if len(result['violation_risks']) > 3:
                                                        st.markdown(f"*...외 {len(result['violation_risks']) - 3}개 추가 위험*")
                                                    st.markdown("---")
                                        else:
                                            st.success("✅ PKL 검색 결과 직접적인 위법 위험은 발견되지 않았습니다.")
                                    else:
                                        st.warning("관련 위법 판례를 찾지 못했습니다.")
                                        
                                except ImportError:
                                    st.error("종합 위법성 분석 모듈을 불러올 수 없습니다.")
                                except Exception as e:
                                    st.error(f"종합 위법성 분석 오류: {str(e)}")
                            
                            # 4-2단계: 기존의 일반적인 PKL 검색
                            # 발견된 문제점을 기반으로 구체적인 검색 쿼리 생성
                            search_terms = []
                            
                            # 사무 관련 문제
                            if any(word in first_analysis for word in ["소관사무", "사무구분", "위임사무", "자치사무"]):
                                search_terms.extend(["기관위임사무 조례제정 불가", "위임사무 조례 제정 한계"])
                            
                            # 법령 위반 관련 문제  
                            if any(word in first_analysis for word in ["법령 위반", "상위법령", "법령우위", "위반"]):
                                search_terms.extend(["법령 위반 조례 사례", "상위법령 충돌 조례"])
                            
                            # 조례 제정 한계 관련
                            if any(word in first_analysis for word in ["제정 한계", "입법한계", "불가", "위법"]):
                                search_terms.extend(["조례 제정 한계 판례", "위법 조례 제정 사례"])
                            
                            # 기본 검색어가 없으면 일반적인 검색어 사용
                            if not search_terms:
                                search_terms = ["법령 위반 조례 판례", "조례 제정 한계 사례"]
                            
                            # 여러 검색어 중 하나 선택 (가장 구체적인 것)
                            search_query_pkl = search_terms[0] if search_terms else "위법 조례 판례"
                            
                            # 향상된 복합 벡터스토어 검색 수행
                            try:
                                from enhanced_search import enhanced_legal_search
                                vectorstore_paths = [
                                    '3. 지방자치단체의 재의·제소 조례 모음집(Ⅸ) (1)_new_vectorstore.pkl'
                                ]
                                enhanced_results = enhanced_legal_search(search_query_pkl, vectorstore_paths, max_results=6)
                                
                                # 기존 형식으로 변환
                                relevant_guidelines = []
                                loaded_stores = set()
                                
                                for result in enhanced_results:
                                    relevant_guidelines.append({
                                        'text': result['text'],
                                        'similarity': result['similarity'],
                                        'distance': 1 - result['similarity'],
                                        'metadata': result['metadata'],
                                        'source_store': result['source_store'].replace('.pkl', '').replace('_', ' ').title(),
                                        'source_file': result['source_store']
                                    })
                                    loaded_stores.add(result['source_store'].replace('.pkl', '').replace('_', ' ').title())
                                
                                loaded_stores = list(loaded_stores)
                                
                            except ImportError:
                                # 기존 방식으로 폴백
                                relevant_guidelines, loaded_stores = search_multiple_vectorstores(
                                    search_query_pkl, 
                                    api_key=gemini_api_key, 
                                    top_k_per_store=2
                                )
                            
                            if relevant_guidelines and loaded_stores:
                                st.success(f"✅ {len(loaded_stores)}개 자료에서 {len(relevant_guidelines)}개 관련 내용을 검색했습니다:")
                                for store in loaded_stores:
                                    st.markdown(f"   • {store}")
                                
                                # 가이드라인 미리보기 (선택사항)
                                with st.expander("📖 검색된 문제 관련 자료 미리보기", expanded=False):
                                    source_groups = {}
                                    for guideline in relevant_guidelines:
                                        source_store = guideline.get('source_store', '알 수 없는 자료')
                                        if source_store not in source_groups:
                                            source_groups[source_store] = []
                                        source_groups[source_store].append(guideline)
                                    
                                    for source_store, guidelines in source_groups.items():
                                        st.markdown(f"**📚 {source_store}**")
                                        for i, guideline in enumerate(guidelines):
                                            similarity_score = guideline.get('similarity', 1-guideline.get('distance', 0))
                                            st.markdown(f"   [{i+1}] (유사도: {similarity_score:.3f})")
                                            st.markdown(guideline['text'][:200] + "..." if len(guideline['text']) > 200 else guideline['text'])
                                            st.markdown("---")
                                
                                # 2차 보강 분석 수행 (조용히)
                                if gemini_api_key:
                                    try:
                                        # 보강 분석용 프롬프트
                                        enhanced_prompt = create_analysis_prompt(
                                            pdf_text,
                                            search_results_for_analysis,
                                            superior_laws_content,
                                            relevant_guidelines,
                                            is_first_ordinance,
                                            comprehensive_analysis_results,
                                            theoretical_results
                                        )
                                        
                                        enhanced_response = model.generate_content(enhanced_prompt)
                                        if enhanced_response and hasattr(enhanced_response, 'text') and enhanced_response.text:
                                            enhanced_analysis = enhanced_response.text
                                            analysis_results.append({
                                                'model': f'Gemini (복합PKL 보강분석 - {len(loaded_stores)}개 자료)',
                                                'content': enhanced_analysis
                                            })
                                    except Exception as e:
                                        st.error(f"복합 PKL 보강 분석 오류: {str(e)}")
                            else:
                                st.info("문제점과 관련된 자료를 찾지 못했습니다.")
                        elif not has_problems:
                            st.info("✅ 문제점이 발견되지 않아 PKL 참고를 건너뜁니다.")
                        elif not use_pkl_auto:
                            st.info("🔄 PKL 자동 참고 기능이 비활성화되어 있습니다.")
                        
                        # 5단계: OpenAI 추가 분석 (선택사항)
                        if openai_api_key:
                            st.info("🔄 5단계: OpenAI 추가 분석을 수행합니다...")
                            try:
                                openai.api_key = openai_api_key
                                # 가장 완전한 프롬프트로 OpenAI 분석
                                openai_prompt = create_analysis_prompt(pdf_text, search_results_for_analysis, superior_laws_content, relevant_guidelines, is_first_ordinance, comprehensive_analysis_results, theoretical_results)
                                
                                response = openai.ChatCompletion.create(
                                    model="gpt-4",
                                    messages=[
                                        {"role": "system", "content": "당신은 법률 전문가입니다. 조례 분석과 검토를 도와주세요."},
                                        {"role": "user", "content": openai_prompt}
                                    ],
                                    temperature=0.7,
                                    max_tokens=4000
                                )
                                
                                if response.choices[0].message.content:
                                    analysis_results.append({
                                        'model': 'OpenAI (추가 분석)',
                                        'content': response.choices[0].message.content
                                    })
                            except Exception as e:
                                st.error(f"OpenAI 분석 오류: {str(e)}")
                                analysis_results.append({
                                    'model': 'OpenAI (추가 분석)',
                                    'error': str(e)
                                })
                        
                        if analysis_results:
                            # 🆕 분석 결과를 세션 상태에 저장
                            st.session_state.analysis_results = analysis_results
                            # 통합 법령 문서 결과도 metadata에 포함
                            final_guideline_results = st.session_state.get('guideline_results', relevant_guidelines)
                            st.session_state.analysis_metadata = {
                                'has_problems': has_problems,
                                'relevant_guidelines': final_guideline_results,
                                'loaded_stores': loaded_stores,
                                'is_first_ordinance': is_first_ordinance,
                                'superior_laws_content': superior_laws_content,
                                'search_results_for_analysis': search_results_for_analysis,
                                'pdf_text': pdf_text,
                                'analysis_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            }

                            # 분석 완료 메시지
                            st.markdown("---")
                            if has_problems and relevant_guidelines and loaded_stores:
                                st.success(f"🎯 **복합 자료 보강 분석 완료**: 문제점 탐지 → {len(loaded_stores)}개 자료 참고 → 보강 분석")
                            elif has_problems and relevant_guidelines:
                                st.success("🎯 **지능형 분석 완료**: 문제점 탐지 → PKL 참고 → 보강 분석")
                            elif has_problems:
                                st.info("⚠️ **문제점 탐지 분석 완료**: PKL 참고 없이 기본 분석만 수행")
                            else:
                                st.success("✅ **기본 분석 완료**: 특별한 문제점이 발견되지 않음")
                            
                            # 분석 결과 요약
                            analysis_count = len([r for r in analysis_results if 'error' not in r])
                            error_count = len([r for r in analysis_results if 'error' in r])
                            
                            if analysis_count > 0:
                                # 🆕 선택된 조례 수 정확히 반영
                                if is_first_ordinance:
                                    analysis_type_text = "최초 제정 조례"
                                else:
                                    selected_count = len(search_results_for_analysis)
                                    analysis_type_text = f"선택된 {selected_count}개 타 시도 조례 비교"
                                st.markdown(f"**📋 분석 유형**: {analysis_type_text}")
                                st.markdown(f"**🤖 수행된 분석**: {analysis_count}개")
                                if relevant_guidelines:
                                    st.markdown(f"**📚 참고된 가이드라인**: {len(relevant_guidelines)}개")
                            
                            # 최종 보고서만 표시 (PKL 보강 분석 또는 OpenAI 분석)
                            final_report = None
                            
                            # 우선순위: 복합PKL 보강분석 > PKL 보강분석 > OpenAI 추가 분석 > 1차 분석
                            for result in reversed(analysis_results):  # 역순으로 최신 결과 우선
                                if 'error' not in result:
                                    if "복합PKL 보강분석" in result['model']:
                                        final_report = result
                                        break
                                    elif "PKL 보강" in result['model'] or "OpenAI" in result['model']:
                                        final_report = result
                                        break
                            
                            # PKL 보강이나 OpenAI가 없으면 1차 분석 사용
                            if not final_report:
                                for result in analysis_results:
                                    if 'error' not in result and "1차 분석" in result['model']:
                                        final_report = result
                                        break
                            
                            # 최종 보고서 표시
                            if final_report:
                                st.markdown("### 📋 최종 분석 보고서")
                                
                                # 보고서 타입 표시
                                if "복합PKL 보강분석" in final_report['model']:
                                    st.success("🎯 **복합 자료 참고 보강 분석 결과**")
                                    st.caption(f"📚 **활용 모델**: {final_report['model']}")
                                elif "PKL 보강" in final_report['model']:
                                    st.success("🎯 **PKL 가이드라인 참고 보강 분석 결과**")
                                elif "OpenAI" in final_report['model']:
                                    st.info("📊 **OpenAI 추가 분석 결과**")
                                else:
                                    st.info("🤖 **Gemini 기본 분석 결과**")
                                
                                # 보고서 내용
                                st.markdown(final_report['content'])
                                
                            # 오류 메시지만 별도 표시
                            for result in analysis_results:
                                if 'error' in result:
                                    st.error(f"❌ {result['model']} 오류: {result['error']}")
                            
                            # Word 문서 생성 및 다운로드
                            with st.spinner("분석 결과 Word 문서 생성 중..."):
                                doc = create_comparison_document(pdf_text, search_results_for_analysis, analysis_results, superior_laws_content, relevant_guidelines)
                                
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_bytes = doc_io.getvalue()
                                
                                # 파일명에 분석 방식 표시
                                if has_problems and relevant_guidelines and loaded_stores:
                                    stores_count = len(loaded_stores)
                                    filename_prefix = f"복합자료보강분석({stores_count}개자료)" if is_first_ordinance else f"조례비교_복합자료분석({stores_count}개자료)"
                                elif has_problems and relevant_guidelines:
                                    filename_prefix = "지능형PKL보강분석" if is_first_ordinance else "조례비교_PKL보강분석"
                                elif has_problems:
                                    filename_prefix = "문제점탐지분석" if is_first_ordinance else "조례비교_문제점분석"
                                else:
                                    filename_prefix = "최초조례_기본분석" if is_first_ordinance else "조례_기본비교분석"
                                
                                st.download_button(
                                    label="📄 분석 결과 Word 문서 다운로드",
                                    data=doc_bytes,
                                    file_name=f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        else:
                            st.error("분석 결과가 없습니다.")

if __name__ == "__main__":
    main()